"""Build WEEK_6_REPORT.docx from the markdown content."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Margins
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Default style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def H1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    p.paragraph_format.space_after = Pt(6)

def H2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)

def H3(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)

def para(text, *, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p

def meta(label, value):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r2 = p.add_run(value)
    p.paragraph_format.space_after = Pt(2)

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    return p

def numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p

def hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BFBFBF")
    pBdr.append(bottom)
    pPr.append(pBdr)

# -------- TITLE --------
H1("MGT 599 Capstone — Weekly Progress Report")

meta("Week", "Week 6 (final week before presentation)")
meta("Reporting period", "May 11 – May 17, 2026")
meta("Group", "4")
meta("Submitted by", "Akash Anipakalu Giridhar")
meta("Submission date", "May 17, 2026")
hr()

# -------- 1. SUMMARY --------
H2("1. Summary")
para(
    "This week the team closed the gap between an inflated headline number and a defensible "
    "honest one. We diagnosed and remediated a data-leakage issue in our earlier evaluation, "
    "rebuilt the train/test pipeline on a company-disjoint basis, retrained our strongest "
    "transformer, and produced a reproducible baseline at 70.29% Macro F1 on the new clean "
    "evaluation. We also locked the foundation for next week's stretch work by shipping six "
    "parallel training variants and finalizing the presentation deck."
)

# -------- 2. QUANTITATIVE OUTCOMES --------
H2("2. Quantitative outcomes")

table = doc.add_table(rows=1, cols=4)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Metric"
hdr[1].text = "Week 5 close"
hdr[2].text = "Week 6 close"
hdr[3].text = "Change"
for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

rows = [
    ("Best reported Macro F1 (Task 1)", "88.90% (later shown to be leaked)", "70.29% (honest, company-disjoint)", "Revised down for integrity"),
    ("ModernBERT-large dev Macro F1 (epoch 3)", "68.28%", "70.29%", "+2.01 pp"),
    ("Industry accuracy", "69.5%", "71.4%", "+1.9 pp"),
    ("Test rows with verified CompanyId", "0 (split lacked the key)", "10,535 / 10,717 (98.3%)", "New artifact"),
    ("Train rows joined to CompanyId", "0", "42,116 / 42,868 (98.2%)", "New artifact"),
    ("Active parallel training runs", "0", "6", "New"),
]
for r in rows:
    row = table.add_row().cells
    for i, v in enumerate(r):
        row[i].text = v

# -------- 3. TASKS COMPLETED --------
H2("3. Tasks completed")

H3("3.1 Data integrity — leakage audit and remediation")
bullet("Audited the cascade pipeline that had been reporting 88.9% Macro F1.")
bullet("Discovered that 97.2% of test rows had been seen by the model during training because the original split was row-level random rather than company-disjoint. The same company's text appeared on both sides of the split.")
bullet("Documented the finding in CASCADE_AUDIT.md with the exact join logic, reproduction steps, and contaminated row counts.")
bullet("Rebuilt the split files using a LongProfile-prefix join (200 characters, with a 100-character fallback) to recover CompanyId on rows that had been stripped of it.")
bullet("Wrote task1_test_with_companyid.csv (10,535 / 10,717 = 98.3% joined) and task1_train_with_companyid.csv (42,116 / 42,868 = 98.2%). These are now the source of truth for every honest evaluation going forward.")

H3("3.2 Modeling — ModernBERT-large retraining on clean splits")
bullet("Re-ran ModernBERT-large training (microsoft/modernbert-large) on the company-disjoint splits.")
bullet("Selected epoch-3 checkpoint based on dev Macro F1 (70.29%); industry accuracy reached 71.4% on the held-out test set.")
bullet("Confirmed the long-tail error profile is concentrated in Diversified Conglomerates (GECS code 31030010), which alone accounts for the largest single contribution to Macro F1 loss.")

H3("3.3 Experimentation — six parallel variants launched")
para("Six Colab notebooks were prepared and queued to explore the configuration space against the new baseline:")
numbered("Baseline ModernBERT-large on raw text (seed 42)")
numbered("Segment-aware text via text_joint field (seed 42)")
numbered("Segment-aware text via text_primary field (seed 42)")
numbered("Segment-aware text with revenue-share sample weighting (seed 42)")
numbered("Distillation on raw text with teacher reasoning JSONL (DISTILL_WEIGHT = 0.3)")
numbered("Variance / ensemble member — segment-aware joint text on seed 123")
para("Each run saves checkpoints, top-5 predictions, and CLS embeddings to Google Drive for downstream ensemble work.")

H3("3.4 Library engineering — BreezeML")
bullet("Continued maintenance of the breezeml PyPI library (Akash A.G., author). The Level 2 hierarchical cascade extension (Sector → Industry Group → Morningstar Code) developed earlier in the term remains the architectural foundation of the V3 Meta-Ensemble.")
bullet("Confirmed five public releases (v0.2.1 through v0.2.5) shipped during the capstone are stable and reproducible.")

H3("3.5 Deliverables for presentation")
bullet("Drafted full 15-slide presentation content matching the Dark Minimalist template palette (#001514 background, #FFFFFF text, #C2D076 lime accent).")
bullet("Wrote PRESENTATION_CONTENT.md with per-slide body, speaker notes, timing breakdown (10-minute target), rubric alignment, and chart generation prompts.")
bullet("Built REFERENCE_DECK.pptx as a layout/color reference for the team.")
bullet("Pre-wrote backup appendix slides (architecture diagram, top confusion class) and an FAQ covering the six most likely panel questions.")

# -------- 4. CHALLENGES --------
H2("4. Challenges encountered")
numbered("Reporting a worse number on purpose. The single largest decision this week was choosing to publish 70.29% honest over 88.9% leaked. It is uncomfortable to walk into a final presentation with a lower headline than we had a month ago, but it is the only defensible position.")
numbered("The 80% target is data-bound, not model-bound. Audit analysis confirmed that 55.2% of training rows have inherent label ambiguity (multi-segment conglomerates with the same LongProfile but different codes per segment). Even a perfect single-code classifier combined with 60% multi-code accuracy mathematically caps Macro F1 at approximately 76%. The remaining headroom requires either an evaluation-frame change (Option A — decidable-subset) or a structural change (Option C — sector-conditioned head on transformer embeddings).")
numbered("Time pressure on parallel runs. Six concurrent training jobs strain Colab Pro Plus session limits. Each notebook is checkpointing to Drive so partial results survive disconnections, but completing all six before the presentation is not guaranteed.")

# -------- 5. PLAN NEXT WEEK --------
H2("5. Plan for next week (Week 7 — post-presentation)")
numbered("Deliver capstone final presentation on Monday, May 18, 2026.")
numbered("Collect the six parallel variant results and select the best two for an ensemble.")
numbered("Begin work on Option C (sector-conditioned head on ModernBERT-large embeddings) — the most promising legitimate path to 75–78% Macro F1.")
numbered("Finalize Task 2 (428 sub-industry code) results and write the supporting documentation.")
numbered("Open-source the company-disjoint split files and the audit script so the methodology is reproducible by future cohorts.")

# -------- 6. REFLECTION --------
H2("6. Reflection")
para(
    "The most valuable thing this week was not a model improvement — it was the audit. "
    "We taught ourselves that an unverified number is worse than no number at all. The honest "
    "70.29% is a lower bar than the leaked 88.9%, but it is the bar from which every future "
    "improvement will be measured truthfully. That discipline, more than any specific architecture "
    "choice, is the work product I am proudest of from this capstone."
)

hr()

# -------- FOOTER --------
meta("Time spent this week", "approximately 38 hours")
meta("Key artifacts produced", "CASCADE_AUDIT.md, task1_test_with_companyid.csv, task1_train_with_companyid.csv, PRESENTATION_CONTENT.md, REFERENCE_DECK.pptx, six configured Colab training notebooks")
meta("Status", "On track for Monday presentation. Stretch work continues in parallel.")

out = r"C:\Users\akash\Desktop\capstone MGT 599\WEEK_6_REPORT.docx"
doc.save(out)
print(f"Saved: {out}")
