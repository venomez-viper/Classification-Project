from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

# ── Helpers ───────────────────────────────────────────────────────────────────
def heading(text, level=1, color=RGBColor(0x1F, 0x35, 0x64)):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = color
    return p

def body(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F2F2F2")
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9.5)
    return p

def divider():
    doc.add_paragraph("─" * 72)

def spacer():
    doc.add_paragraph("")

# ── Cover ─────────────────────────────────────────────────────────────────────
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("MGT 599 Capstone — Group 4")
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run("Week 3 — Classifier Testing Assignments")
r2.font.size = Pt(13)
r2.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run("Strayer University  |  April 2026")
r3.font.size = Pt(11)
r3.italic = True
r3.font.color.rgb = RGBColor(0x90, 0x90, 0x90)

spacer()
divider()
spacer()

# ── Context ───────────────────────────────────────────────────────────────────
heading("Overview", level=1)
body(
    "The core pipeline (data cleaning, feature engineering, and the baseline Linear SVM) "
    "is already complete. Your job this week is to run a different classifier on both tasks "
    "using the breezeml package, record the results, and write a short interpretation. "
    "The team lead will compile all results into the final comparison table.",
    size=11,
)
spacer()
body("Two classification tasks:", bold=True)
bullet("Task 1 — Predict broad industry (MstarGlobal) from company long-form description")
bullet("Task 2 — Predict subindustry from business segment description")
spacer()
body(
    "Each member below is assigned one classifier. Run it on BOTH tasks following the "
    "template provided. Do not change the data loading or vectorizer settings — keep "
    "them identical to the baseline so results are comparable.",
    size=11,
)

spacer()
divider()
spacer()

# ── Setup section ─────────────────────────────────────────────────────────────
heading("Before You Start — Setup Checklist", level=1)
bullet("Pull the latest main branch:  git pull origin main")
bullet("Confirm  data/cleaned/task1_clean.csv  and  task2_clean.csv  exist (ask team lead if missing)")
bullet("Install dependencies if needed:")
code_block("pip install pandas scikit-learn breezeml")
bullet("Create your own branch before making any changes:")
code_block("git checkout -b member<N>-classifier")
spacer()

divider()
spacer()

# ── Member assignments ─────────────────────────────────────────────────────────
members = [
    {
        "num": 1,
        "classifier": "Logistic Regression",
        "fn": "classifiers.logistic",
        "call_t1": "model, report = classifiers.logistic(X=X_vec, y=y)",
        "call_t2": "model, report = classifiers.logistic(X=X_vec, y=y)",
        "note": (
            "Logistic Regression is the standard linear baseline for text classification. "
            "It is fast and interpretable. Compare its Macro F1 to the Linear SVM — they "
            "should be in the same ballpark. If Logistic is notably lower, note it."
        ),
    },
    {
        "num": 2,
        "classifier": "Multinomial Naive Bayes",
        "fn": "classifiers.multinomial_nb",
        "call_t1": "model, report = classifiers.multinomial_nb(X=X_vec, y=y)",
        "call_t2": "model, report = classifiers.multinomial_nb(X=X_vec, y=y)",
        "note": (
            "Multinomial Naive Bayes is designed for word-count / TF-IDF features and is "
            "extremely fast. It often performs well on text data despite its simplicity. "
            "Pay attention to whether it struggles with the large number of subindustry "
            "classes in Task 2."
        ),
    },
    {
        "num": 3,
        "classifier": "Decision Tree",
        "fn": "classifiers.decision_tree",
        "call_t1": "model, report = classifiers.decision_tree(X=X_vec, y=y)",
        "call_t2": "model, report = classifiers.decision_tree(X=X_vec, y=y)",
        "note": (
            "Decision Trees are interpretable but tend to overfit on high-dimensional TF-IDF "
            "features. Expect lower F1 scores compared to SVM or Logistic Regression. "
            "That is useful information — note the gap and explain why it likely occurs "
            "(high dimensionality, sparse features)."
        ),
    },
    {
        "num": 4,
        "classifier": "Extra Trees",
        "fn": "classifiers.extra_trees",
        "call_t1": "model, report = classifiers.extra_trees(X=X_vec, y=y)",
        "call_t2": "model, report = classifiers.extra_trees(X=X_vec, y=y)",
        "note": (
            "Extra Trees (Extremely Randomized Trees) is an ensemble method similar to "
            "Random Forest but with more randomness in splits, which makes it faster to train. "
            "Compare its results to the Random Forest baseline already in the models/ folder. "
            "It may perform similarly or slightly differently — document the difference."
        ),
    },
]

for m in members:
    heading(f"Member {m['num']} — {m['classifier']}", level=1)

    body("Assigned Classifier:", bold=True)
    body(f"  {m['fn']}()", size=11)
    spacer()

    body("Interpretation note:", bold=True)
    body(f"  {m['note']}", size=11)
    spacer()

    # ── Task 1 template ───────────────────────────────────────────────────────
    heading("Task 1 — Broad Industry Classification (MstarGlobal)", level=2)
    body("Copy this code into a new notebook or script. Run it top to bottom.", italic=True, size=10)
    spacer()

    code_block(
        "import pandas as pd\n"
        "from sklearn.feature_extraction.text import TfidfVectorizer\n"
        "from sklearn.model_selection import train_test_split\n"
        "from sklearn.metrics import f1_score, classification_report\n"
        "from breezeml import classifiers\n"
        "\n"
        "# 1. Load data\n"
        "df = pd.read_csv('../data/cleaned/task1_clean.csv')\n"
        "df['Combined_Text'] = (\n"
        "    df['LongProfile'].fillna('') + ' ' +\n"
        "    df['SegmentName'].fillna('') + ' ' +\n"
        "    df['SegmentDescription'].fillna('')\n"
        ")\n"
        "\n"
        "# 2. Vectorize  — keep identical to baseline\n"
        "vectorizer = TfidfVectorizer(\n"
        "    max_features=50000, sublinear_tf=True,\n"
        "    stop_words='english', ngram_range=(1, 2)\n"
        ")\n"
        "X_vec = vectorizer.fit_transform(df['Combined_Text'])\n"
        "y = df['MstarGlobal'].values\n"
        "\n"
        "# 3. Train/test split\n"
        "X_tr, X_te, y_tr, y_te = train_test_split(\n"
        "    X_vec, y, test_size=0.2, random_state=42, stratify=y\n"
        ")\n"
        "\n"
        f"# 4. Train with breezeml\n"
        f"{m['call_t1'].replace('X=X_vec, y=y', 'X=X_tr, y=y_tr')}\n"
        "\n"
        "# 5. Evaluate\n"
        "y_pred = model.predict(X_te)\n"
        "macro_f1    = f1_score(y_te, y_pred, average='macro')\n"
        "weighted_f1 = f1_score(y_te, y_pred, average='weighted')\n"
        "print(f'MACRO F1    : {macro_f1 * 100:.2f}%')\n"
        "print(f'WEIGHTED F1 : {weighted_f1 * 100:.2f}%')\n"
        "print(classification_report(y_te, y_pred, zero_division=0))"
    )

    spacer()
    body("Record your results here:", bold=True)
    bullet("Macro F1 Score (Task 1):  __________%")
    bullet("Weighted F1 Score (Task 1):  __________%")
    bullet("Training time (approximate):  __________")
    spacer()

    # ── Task 2 template ───────────────────────────────────────────────────────
    heading("Task 2 — Subindustry Classification", level=2)
    body("Same process, different dataset and target column.", italic=True, size=10)
    spacer()

    code_block(
        "import pandas as pd\n"
        "from sklearn.feature_extraction.text import TfidfVectorizer\n"
        "from sklearn.model_selection import train_test_split\n"
        "from sklearn.metrics import f1_score, classification_report\n"
        "from breezeml import classifiers\n"
        "\n"
        "# 1. Load data\n"
        "df = pd.read_csv('../data/cleaned/task2_clean.csv')\n"
        "\n"
        "# Filter classes with fewer than 5 samples\n"
        "counts = df['Subindustry'].value_counts()\n"
        "valid  = counts[counts > 5].index\n"
        "df     = df[df['Subindustry'].isin(valid)]\n"
        "\n"
        "df['Combined_Text'] = (\n"
        "    df['SegmentName'].fillna('') + ' ' +\n"
        "    df['SegmentDescription'].fillna('')\n"
        ")\n"
        "\n"
        "# 2. Vectorize  — keep identical to baseline\n"
        "vectorizer = TfidfVectorizer(\n"
        "    max_features=10000, sublinear_tf=True,\n"
        "    stop_words='english', ngram_range=(1, 2)\n"
        ")\n"
        "X_vec = vectorizer.fit_transform(df['Combined_Text'])\n"
        "y = df['Subindustry'].values\n"
        "\n"
        "# 3. Train/test split\n"
        "X_tr, X_te, y_tr, y_te = train_test_split(\n"
        "    X_vec, y, test_size=0.2, random_state=42, stratify=y\n"
        ")\n"
        "\n"
        f"# 4. Train with breezeml\n"
        f"{m['call_t2'].replace('X=X_vec, y=y', 'X=X_tr, y=y_tr')}\n"
        "\n"
        "# 5. Evaluate\n"
        "y_pred = model.predict(X_te)\n"
        "macro_f1    = f1_score(y_te, y_pred, average='macro')\n"
        "weighted_f1 = f1_score(y_te, y_pred, average='weighted')\n"
        "print(f'MACRO F1    : {macro_f1 * 100:.2f}%')\n"
        "print(f'WEIGHTED F1 : {weighted_f1 * 100:.2f}%')\n"
        "print(classification_report(y_te, y_pred, zero_division=0))"
    )

    spacer()
    body("Record your results here:", bold=True)
    bullet("Macro F1 Score (Task 2):  __________%")
    bullet("Weighted F1 Score (Task 2):  __________%")
    bullet("Training time (approximate):  __________")
    spacer()

    # ── Deliverables ──────────────────────────────────────────────────────────
    heading("Deliverables", level=2)
    bullet(
        "Your notebook or script file saved as  "
        f"notebooks/week3_member{m['num']}_{m['classifier'].lower().replace(' ', '_')}.ipynb"
    )
    bullet(
        "A handoff note filled in at  "
        f"docs/handoff_notes/member{m['num']}_handoff.md  "
        "(use the template already in that folder)"
    )
    bullet("Post your two Macro F1 scores in the group chat so the team lead can build the comparison table")
    spacer()

    if m["num"] < 4:
        divider()
        spacer()

# ── Footer ────────────────────────────────────────────────────────────────────
divider()
spacer()
heading("Deadline and Questions", level=1)
body(
    "Submit your notebook and handoff note before the Week 3 deadline. "
    "If you hit an import error, make sure breezeml is installed in the same "
    "Python environment you are running the notebook from. If you hit a data "
    "error, confirm that main.py has been run and the cleaned CSVs exist. "
    "For anything else, message the team lead.",
    size=11,
)

spacer()
body("MGT 599 Capstone — Group 4  |  Strayer University  |  April 2026", italic=True, size=10)

# ── Save ──────────────────────────────────────────────────────────────────────
out = r"C:\Users\akash\Desktop\capstone MGT 599\Week3_Team_Classifier_Assignments.docx"
doc.save(out)
print(f"Saved: {out}")
