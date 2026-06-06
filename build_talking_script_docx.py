"""Build TALKING_SCRIPT.docx — printable, rehearsal-ready, 5 speakers."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
MUTED = RGBColor(0x6B, 0x6B, 0x6B)
DARK = RGBColor(0x1A, 0x1A, 0x1A)

def H1(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(4)

def H2(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(15); r.font.color.rgb = ACCENT
    p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(4)

def H3(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = DARK
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)

def para(text, italic=False, bold=False, size=11, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text); r.italic = italic; r.bold = bold; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return p

def meta(label, value):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: "); r1.bold = True
    r2 = p.add_run(value)
    p.paragraph_format.space_after = Pt(2)

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p

def checkbox(text):
    p = doc.add_paragraph()
    r1 = p.add_run("☐  "); r1.font.size = Pt(12)
    r2 = p.add_run(text)
    return p

def script_block(speaker, time_str):
    p = doc.add_paragraph()
    r1 = p.add_run(f"SPEAKER: {speaker}"); r1.bold = True; r1.font.color.rgb = ACCENT; r1.font.size = Pt(12)
    r2 = p.add_run(f"     TIME: {time_str}"); r2.bold = True; r2.font.size = Pt(11); r2.font.color.rgb = MUTED
    p.paragraph_format.space_after = Pt(2)

def quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.font.size = Pt(12); r.font.name = "Cambria"
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "12"); left.set(qn("w:space"), "8"); left.set(qn("w:color"), "1F4E79")
    pBdr.append(left); pPr.append(pBdr)

def handoff(text):
    p = doc.add_paragraph()
    r = p.add_run(f"→ {text}")
    r.italic = True; r.font.size = Pt(10); r.font.color.rgb = MUTED
    p.paragraph_format.space_after = Pt(10)

def hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6"); bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "BFBFBF")
    pBdr.append(bottom); pPr.append(pBdr)

def page_break():
    p = doc.add_paragraph()
    r = p.add_run(); r.add_break(WD_BREAK.PAGE)

# ============== HEADER ==============
H1("Mapping Market Reality — Team Talking Script")
meta("Event", "MGT 599 Capstone Final Presentation")
meta("Group", "4 · DePaul University · Kellstadt Graduate School of Business")
meta("Date", "May 18, 2026")
meta("Deck", "Mapping_Market_Reality.pptx (12 slides)")
meta("Target runtime", "10:00 + Q&A")
hr()

# ============== TEAM ==============
H2("The team")
para("Akash Anipakalu Giridhar  ·  Subasree Segar  ·  Tserennadmid Batkhuu  ·  Srilaxmi Ganjipalli  ·  Vishal Shaileshkumar Rathod", bold=True)

# ============== SPEAKER ASSIGNMENTS TABLE ==============
H2("Speaker assignments (5 speakers)")
tbl = doc.add_table(rows=1, cols=4)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
for i, t in enumerate(["Speaker", "Role / Voice", "Slides", "Time"]):
    hdr[i].text = t
    for p in hdr[i].paragraphs:
        for r in p.runs: r.bold = True
rows = [
    ("Tserennadmid Batkhuu", "Host — opens and closes the deck, owns continuity", "1, 2, 12", "1:20"),
    ("Srilaxmi Ganjipalli", "Business framing — problem statement + closing recap", "3, 11", "1:30"),
    ("Akash Anipakalu Giridhar", "Architecture, BreezeML, ModernBERT, path forward", "4, 6, 10", "2:45"),
    ("Vishal Shaileshkumar Rathod", "Data foundation, feature engineering, models", "5, 7", "1:50"),
    ("Subasree Segar", "Evaluation — the audit and honest baseline (centerpiece)", "8, 9", "2:45"),
]
for r in rows:
    cells = tbl.add_row().cells
    for i, v in enumerate(r): cells[i].text = v

# ============== HOUSE RULES ==============
H2("House rules for delivery")
bullet("Read your slides aloud twice tonight, with a stopwatch. Adjust pace, not content.")
bullet("Look at the panel, not the screen. The slide is behind you; trust it.")
bullet("Hand off explicitly using the cue printed at the bottom of each slide.")
bullet("Pause for one full breath before any number ending in a percent sign — numbers are the strongest moments.")
bullet("If you forget a line, stop. Look up. Continue. Silence is not failure.")
bullet("Recommend: Tserennadmid holds the clicker (opens and closes; no mid-talk clicker handoff).")

page_break()

# ============== SLIDES ==============
slides_data = [
    (1, "Tserennadmid", "0:30", "Title — Mapping Market Reality",
     "Good afternoon. Our team is Group 4, and our capstone asks a single question — can a machine read a company description and know what industry it belongs to? That sounds simple. It is not. Today we will walk you through six months of work answering it: the problem we set out to solve, the data foundation we built, the moment our results turned out to be too good to be true, and the honest baseline we now stand on. I am Tserennadmid. Joining me are Akash, Subasree, Srilaxmi, and Vishal.",
     "Tserennadmid advances to Slide 2 (Agenda)."),

    (2, "Tserennadmid", "0:30", "Agenda",
     "Here is how we will move. Srilaxmi will set up the problem. Akash will walk you through the pipeline we built, including the open-source library we authored along the way. Vishal will show you the data foundation and the four models we trained. Subasree will take you through the audit moment that reshaped this project. Akash returns with the path forward, Srilaxmi closes with what we built and learned, and I will wrap with the takeaways.",
     "Hand off: Tserennadmid → Srilaxmi."),

    (3, "Srilaxmi", "0:45", "The Problem",
     "Thank you Tserennadmid. Morningstar classifies every public company on Earth into one of 145 industries and 428 sub-industries. In our dataset, that is 23,207 unique companies spread across 53,585 segment-level records. Right now, that classification happens by hand. Analysts read each company's profile and assign a code. It is slow, it is inconsistent between analysts, and it does not scale to IPOs, spinoffs, or mergers. And it has real money attached — sector ETFs, peer benchmarks, factor models all roll up from these codes. A misclassified conglomerate distorts the index. Akash will explain how we built a system to do this work automatically.",
     "Hand off: Srilaxmi → Akash."),

    (4, "Akash", "0:45", "The Pipeline + Team",
     "Thanks Srilaxmi. Three principles guided how we built this. First — hierarchy-first. We do not ask the model to choose between 145 classes at once. We route Sector to Industry Group to Code, the way an analyst actually thinks. Second — audit-first. Every number you will see today is backed by a strict company-disjoint test split. Third — taxonomy-grounded. Our predictions are anchored to Morningstar's 2019 GECS definitions, not to whatever the model decides to invent. On the team side, I own architecture, the BreezeML library, and ModernBERT fine-tuning. Subasree owns evaluation. Vishal owns feature engineering. Tserennadmid owns the data pipeline and repo. Srilaxmi anchors the business framing. Vishal will tell you what the data looks like.",
     "Hand off: Akash → Vishal."),

    (5, "Vishal", "0:50", "Data Foundation",
     "Thanks Akash. The raw input to our system is a company's LongProfile — a free-form, unstructured business description. That alone is not enough, because of one fact that drove most of our engineering decisions: 35.1 percent of companies in our dataset are multi-segment conglomerates. One company, multiple legitimate industry codes. That means 55.2 percent of our database rows have inherent label ambiguity built in. So we built an enrichment layer on top of the LongProfile — adding SegmentName, SegmentDescription, Revenue, and revenue share. This gives the model what an analyst sees: not just what the company is, but which segment actually makes the money. Akash will show you the library we built to make this possible.",
     "Hand off: Vishal → Akash."),

    (6, "Akash", "1:15", "BreezeML ★",
     "Most capstone teams use libraries. We shipped one. [pause] This is breezeml — pip install breezeml. It is a zero-boilerplate machine learning framework I authored on top of scikit-learn, and it is live on PyPI right now. Anyone in this room can install it. We did not just write it once. We shipped five public releases during this project. The most important one — version 0.2.3 — was a primal SVM fix. The library was deadlocking on our text data because of how scikit-learn solved the support vector formulation. Most teams would have switched models. We diagnosed the math, patched the upstream library, and brought training time from over twenty minutes down to under two seconds. Version 0.2.5 added balanced class weights for the rare GECS labels, which eliminated the need for SMOTE oversampling entirely. The result is a public library that any future analyst can install — including the panel — and replicate our preprocessing in one line. Vishal will walk you through what we did with it.",
     "Hand off: Akash → Vishal."),

    (7, "Vishal", "1:00", "Four Models, Four Steps",
     "We climbed this wall in four steps. Our TF-IDF cascade with LinearSVC got us to 59.65 percent macro F1. Stacking twelve classical models in our V8 mega-ensemble took us to 68.42. We then moved to transformers — ModernBERT-base v2 landed at 67.18 percent. And our current honest best, ModernBERT-large at epoch three, is 70.29 percent macro F1. The target — the gold line you see across the chart — is 80 percent. [pause] Notice the shape of the climb. Each jump is smaller than the last. By the time we hit ModernBERT-large, three hours of GPU time was buying us a single point of macro F1. Something else was happening — and Subasree will tell you what we found.",
     "Hand off: Vishal → Subasree."),

    (8, "Subasree", "1:30", "The Mirage ★ — centerpiece, slow down",
     "Thank you Vishal. Earlier in this project, our cascade reported 88.9 percent macro F1. That is the number on this slide — the one with the red strikethrough. [pause] We did not trust it. So we audited the data. And what we found was this: 97.2 percent of our test rows had already been seen by the model during training. The same company's text appeared on both sides of the train-test split, because the original split was randomized at the row level, not at the company level. The 88.9 was not a result. It was a leak. [pause] Three things happened next. We acknowledged it. We documented it in a file called CASCADE_AUDIT.md, with the exact join logic and contaminated row counts. And we rebuilt the entire evaluation pipeline on a strict company-disjoint basis — meaning no company ever appears in both train and test again. We could have published the 88.9. Nobody outside this team would have caught it. That decision to report a lower honest number is, in my view, the most professional thing this team did in the entire capstone.",
     "Subasree advances to Slide 9."),

    (9, "Subasree", "1:15", "Honest Baseline + Ceiling",
     "Here is where we actually stand. ModernBERT-large fine-tuned on the company-disjoint splits delivers 70.29 percent dev macro F1. That number is reproducible, it is auditable, and it is the bar from which everything we do next will be measured. But notice the chart on the right. Our test set is 45.3 percent single-code companies, 25.5 percent two-code, 15.9 percent three-code, and 13.3 percent four-or-more-code conglomerates. Even a hypothetical perfect classifier on the single-code rows, combined with 60 percent accuracy on the multi-code rows, mathematically caps macro F1 at approximately 76 percent. [pause] In other words: the gap from 70 to 80 percent is not a model problem. It is a data problem. The ceiling is structural. Akash will show you the four strategies we are using to attack it.",
     "Hand off: Subasree → Akash."),

    (10, "Akash", "0:45", "Four Paths",
     "We have four paths forward. Path A — the decidable-subset approach — scores models only on rows with one unambiguous label, bypassing the conglomerate ceiling. Path B rolls segment-level predictions up to company-level via revenue weighting. Path D extends ModernBERT training with longer schedules — but the data confirms that route caps near 73 percent. [pause and point to Path C] The path we believe in is C. We apply a hierarchical routing head directly onto ModernBERT-large embeddings — a fusion of the cascade architecture and the transformer's representational power. This run is in flight right now, and our projection is 75 to 78 percent macro F1 by term end. Srilaxmi will tell you what this all adds up to.",
     "Hand off: Akash → Srilaxmi."),

    (11, "Srilaxmi", "0:45", "What We Built · What We Learned",
     "Thanks Akash. Here is what we built. A ModernBERT-large classifier running at an honest 70.29 percent macro F1. The breezeml library, with five successful PyPI releases. Task 2 — the 428 sub-industry models — successfully trained across all distinct codes. And a production-ready backend with a unified Next.js dashboard. [pause] Here is what we learned. Audit your own numbers, relentlessly, before anyone else does. The structural data ceiling will beat the mathematical model every single time. And a hierarchy-first architecture, consistently, defeats end-to-end processing on long-tail classification problems. [pause] If we leave you with one sentence, it is this: industry classification is not one hundred percent automatable. The resulting product is an analyst-first system, designed to handle the obvious cases, defer on the ambiguous conglomerates, and log every prediction against the truth.",
     "Hand off: Srilaxmi → Tserennadmid."),

    (12, "Tserennadmid", "0:20", "Thank You & Questions",
     "Thank you Srilaxmi. Three takeaways to remember. Data quality shapes model truth. Strict company-disjoint splits reveal reality. And Path C — the hierarchical routing head — is the best legitimate path forward. [pause] On behalf of Akash, Subasree, Srilaxmi, Vishal, and myself — thank you. We are happy to take your questions.",
     "Open the floor for Q&A."),
]

for idx, speaker, t, title, words, ho in slides_data:
    H2(f"Slide {idx} — {title}")
    script_block(speaker, t)
    quote(words)
    handoff(ho)

page_break()

# ============== RUNTIME TABLE ==============
H2("Runtime check")
rt = doc.add_table(rows=1, cols=5)
rt.style = "Light Grid Accent 1"
hdr = rt.rows[0].cells
for i, t in enumerate(["Slide", "Speaker", "Target", "Words", "Pace"]):
    hdr[i].text = t
    for p in hdr[i].paragraphs:
        for r in p.runs: r.bold = True
runtime_rows = [
    ("1", "Tserennadmid", "0:30", "~85", "comfortable"),
    ("2", "Tserennadmid", "0:30", "~90", "comfortable"),
    ("3", "Srilaxmi", "0:45", "~115", "comfortable"),
    ("4", "Akash", "0:45", "~135", "brisk"),
    ("5", "Vishal", "0:50", "~135", "brisk"),
    ("6", "Akash", "1:15", "~225", "normal"),
    ("7", "Vishal", "1:00", "~155", "normal"),
    ("8", "Subasree", "1:30", "~255", "normal — slow down"),
    ("9", "Subasree", "1:15", "~195", "normal"),
    ("10", "Akash", "0:45", "~130", "brisk"),
    ("11", "Srilaxmi", "0:45", "~165", "brisk"),
    ("12", "Tserennadmid", "0:20", "~55", "comfortable"),
]
for r in runtime_rows:
    cells = rt.add_row().cells
    for i, v in enumerate(r): cells[i].text = v
para("Total: approximately 10:10. Aim to land at 10:00 by trimming one or two pauses.", italic=True)

# ============== Q&A ==============
H2("Q&A prep — six likely questions")

qa = [
    ("Why did the 88.9 percent number drop to 70?", "Subasree",
     "The 88.9 was a leakage artifact. Our original splits were row-randomized, which let the same company's text appear in both train and test. We audited it, documented it in CASCADE_AUDIT.md, rebuilt the splits on a company-disjoint basis, and re-baselined. The 70.29 is reproducible on splits where no company appears in both halves."),
    ("Why ModernBERT and not GPT-4 or Claude?", "Akash",
     "Our project is closed-stack — no external APIs were permitted for fair evaluation. ModernBERT-large also outperformed DeBERTa on our dev set by 2.1 points and handles the full LongProfile in its 8K context window without truncation. It runs on a single GPU in production."),
    ("What is your validation strategy?", "Subasree",
     "Strict company-disjoint splits — CompanyId is the join key, and no company appears in both train and test. We hold out 10,535 of 10,717 test rows after CompanyId joining, which is 98.3 percent coverage. We report macro F1 on dev for model selection and macro F1 on the held-out test for final numbers."),
    ("Is BreezeML actually used outside this project?", "Akash",
     "It is public on PyPI. Adoption is modest — this is a capstone, not a startup — but the patches are mathematically correct for any high-dimensional text classification task, so they are useful beyond our scope."),
    ("Why are you confident in 75 to 78 percent for Path C?", "Akash",
     "Path C uses ModernBERT-large's embeddings as input to a hierarchical routing head trained on the GECS taxonomy. The 70.29 is what the model achieves when forced to choose among 145 classes flat. The hierarchical routing collapses the long-tail problem into successive smaller decisions, which is mathematically equivalent to giving the model the taxonomy structure as a prior. Empirical results from related work suggest 5 to 8 point gains."),
    ("How is this different from what Morningstar already does internally?", "Srilaxmi or Vishal",
     "Morningstar's process is analyst-driven, slow, and inconsistent across humans. Our system is reproducible, sub-second per company, and audited end-to-end. It does not replace the analyst — it gives them a defensible starting point with calibrated uncertainty, so they spend their time on the 35 percent of cases where ambiguity is real."),
]

for q, who, a in qa:
    H3(f"Q: {q}")
    p = doc.add_paragraph()
    r1 = p.add_run(f"Answer (lead: {who}): "); r1.bold = True; r1.font.color.rgb = ACCENT
    r2 = p.add_run(a)
    p.paragraph_format.space_after = Pt(6)

# ============== REHEARSAL ==============
H2("Rehearsal checklist (do before walking in)")
checkbox("Each speaker reads their slides aloud at least twice, with a stopwatch")
checkbox("Full team walkthrough once, end-to-end, with handoff phrases practiced")
checkbox("Confirm the deck file opens on the presentation machine")
checkbox("Print this script — one copy per speaker — as backup if a laptop fails")
checkbox("Tserennadmid holds the clicker the whole time")
checkbox("Open team_briefing.html or the demo in a browser tab beforehand, in case of a live-demo question")
checkbox("Bring water. Talk slow. Smile once.")

out = r"C:\Users\akash\Desktop\capstone MGT 599\TALKING_SCRIPT_v2.docx"
doc.save(out)
print(f"Saved: {out}")
