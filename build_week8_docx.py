"""Build WEEK_8_REPORT.docx from the markdown content."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def H1(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor(0x1A,0x1A,0x1A)
    p.paragraph_format.space_after = Pt(6)

def H2(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x1A,0x1A,0x1A)
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)

def H3(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)

def para(text):
    p = doc.add_paragraph(); p.add_run(text); return p

def meta(label, value):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: "); r1.bold = True
    p.add_run(value)
    p.paragraph_format.space_after = Pt(2)

def bullet(text):
    p = doc.add_paragraph(style="List Bullet"); p.add_run(text); return p

def numbered(text):
    p = doc.add_paragraph(style="List Number"); p.add_run(text); return p

def hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "BFBFBF")
    pBdr.append(bottom); pPr.append(pBdr)

# -------- TITLE --------
H1("MGT 599 Capstone — Weekly Progress Report")
meta("Week", "Week 8 (deployment, Task 2 completion, site audit)")
meta("Reporting period", "May 25 – May 31, 2026")
meta("Group", "4")
meta("Submitted by", "Akash Anipakalu Giridhar")
meta("Submission date", "May 31, 2026")
hr()

# -------- 1. SUMMARY --------
H2("1. Summary")
para(
    "This week closed two open threads from Week 7 and converted the project from a local research prototype "
    "into a live, publicly accessible product. Task 2 (sub-industry classification, 428 classes) was completed "
    "at 55.44% Macro F1 using a constrained L4 cascade that enforces the GECS parent-child hierarchy at "
    "inference time. The full production stack was deployed: the cascade SVM runs on a public Hugging Face "
    "Space, the ModernBERT-large demo runs on a second Space, and the Next.js frontend is live on Vercel. "
    "A content audit of the frontend identified and corrected stale numbers, wrong class counts, outdated "
    "architecture descriptions, and incorrect deployment references throughout the site."
)

# -------- 2. QUANTITATIVE OUTCOMES --------
H2("2. Quantitative outcomes")
table = doc.add_table(rows=1, cols=4)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
for i, t in enumerate(["Metric", "Week 7 close", "Week 8 close", "Change"]):
    hdr[i].text = t
    for p in hdr[i].paragraphs:
        for r in p.runs: r.bold = True
rows = [
    ("Task 1 Macro F1 (locked)",        "75.0%",    "75.0%",                  "— (locked)"),
    ("Task 1 top-3 accuracy",            "91.4%",    "91.4%",                  "— (locked)"),
    ("Task 2 Macro F1",                  "not measured", "55.44%",             "new"),
    ("Task 2 sub-industry classes",      "—",        "428",                    "confirmed"),
    ("Production deployment",            "none",     "Vercel + 2x HF Space",   "new"),
    ("Site pages with stale content",    "~12",      "0",                      "audited"),
    ("Build errors resolved",            "—",        "6",                      "new"),
]
for r in rows:
    row = table.add_row().cells
    for i, v in enumerate(r): row[i].text = v

# -------- 3. TASKS COMPLETED --------
H2("3. Tasks completed")

H3("3.1 Task 2 constrained L4 cascade — 55.44% Macro F1")
para(
    "The Task 2 prediction pipeline was completed using a constrained 4-level cascade architecture. The L4 "
    "sub-industry classifier is conditioned on the Task 1 output: for each company, the prediction space is "
    "restricted to the valid GECS children of the predicted Task 1 industry code via a task1_to_task2_map "
    "lookup, reducing the effective search space from 428 classes to approximately 3–15 per row."
)
para("Model components:")
bullet("Vectorizer: t2_cascade_seg_vec.pkl — segment-aware TF-IDF trained on Task 2 rows only")
bullet("Classifier: t2_cascade_L4_seg.joblib — LinearSVC per Task 1 parent group")
bullet("Constraint: Task 1 predicted code enforces valid sub-industry candidates before L4 ranking")
bullet("Result: 55.44% Macro F1 across 428 sub-industry codes")
para(
    "This result is lower than Task 1 for structural reasons: 65% of the 428 sub-industry classes have fewer "
    "than 10 training samples, and the hierarchical constraint means Task 1 errors propagate — a wrong sector "
    "prediction at L1 makes the correct sub-industry mathematically unreachable. The 55.44% number is reported "
    "without adjustment or scope restriction; it is the honest result on the full 428-class test distribution."
)

H3("3.2 Production deployment — Hugging Face Spaces")
para("Two public Hugging Face Spaces were deployed:")
para("Space 1 — GECS Cascade SVM (akash-ag-gecs-classifier-space.hf.space)")
bullet("Serves the Task 1 + Task 2 cascade at POST /api/predict")
bullet("FastAPI + Gradio (Docker SDK), hf_space/app.py")
bullet("Returns mstar_code, mstar_label, confidence_t1, alternatives_t1, sub_code, sub_label, confidence_t2, alternatives_t2, route_reason")
para("Space 2 — ModernBERT Demo (akash-ag-gecs-modernbert.hf.space)")
bullet("Serves the ModernBERT-large inference demo at POST /api/predict")
bullet("Separate space to isolate GPU cold-start behavior from the SVM demo")
para(
    "A key operational finding: HF Spaces sleep after inactivity, with cold-start latency of 30–60 seconds. "
    "All Vercel proxy routes were given export const maxDuration = 60 to prevent the 10-second default timeout "
    "from killing requests during space warm-up."
)

H3("3.3 Vercel frontend deployment")
para(
    "The Next.js 15 frontend was deployed to Vercel with all five API routes proxying to the HF Spaces. "
    "Six build-blocking errors were diagnosed and resolved in sequence:"
)
numbered(
    "Missing packages — app/ui-demo/page.tsx imported @radix-ui/react-accordion, @radix-ui/react-icons, "
    "and six @visx/* packages not in package.json. Resolution: removed the unused demo sandbox page."
)
numbered(
    "SSR prerender crash on /journey — RevealText uses framer-motion spring+scale transforms (type: 'spring', "
    "scale: 0 to 1). framer-motion v12 cannot resolve the property type registry during Next.js static "
    "prerendering. Resolution: 'use client' + next/dynamic with ssr: false on the journey and team pages."
)
numbered(
    "SSR prerender crash on / — FallingPattern (used in Hero) animates backgroundPosition with multi-value "
    "strings. framer-motion v12 throws 'Cannot destructure property base of f[c] as it is undefined' when "
    "parsing this property during server render. Resolution: same dynamic/ssr:false pattern on app/page.tsx."
)
numbered(
    "ssr: false not allowed in server components — Next.js App Router requires the calling component to be "
    "a client component before ssr: false is valid. Resolution: added 'use client' to the three affected page files."
)
numbered(
    "API routes calling dead paths — predict/route.ts called /api/predict_legendary, predict_llm/route.ts "
    "called /predict, predict_legendary/route.ts and predict_routed/route.ts called /predict. The HF platform "
    "proxy only forwards /api/predict to the Space container. All four routes updated to call /api/predict."
)
numbered(
    "HF Space returning 404 for all paths — The Space had been created as private. All server-side fetch calls "
    "from Vercel (unauthenticated) hit the HF platform 404 page. Resolution: changed Space visibility to "
    "public in HF settings."
)

H3("3.4 Frontend content audit and corrections")
para(
    "A systematic pass across all dashboard tabs and pages identified and corrected the following stale content:"
)

table2 = doc.add_table(rows=1, cols=3)
table2.style = "Light Grid Accent 1"
hdr2 = table2.rows[0].cells
for i, t in enumerate(["Location", "Error", "Correction"]):
    hdr2[i].text = t
    for p in hdr2[i].paragraphs:
        for r in p.runs: r.bold = True
content_rows = [
    ("All tabs", "407 classes (Task 2)", "428 classes throughout"),
    ("ModelsTab.tsx", "DeBERTa-v3-small as Experimental Track", "ModernBERT-large ensemble (75.0%)"),
    ("MonitoringTab.tsx", "F1 series at 86-87% (leaked number)", "74.5-75.5% (honest number)"),
    ("DeploymentTab.tsx", "Port 5000 ACTIVE", "HF SPACE ACTIVE"),
    ("CodeTab.tsx", "Flask @app.route on localhost", "FastAPI hf_space/app.py cloud code"),
    ("DocumentationTab.tsx", "80/20 stratified split", "Company-disjoint split"),
    ("DocumentationTab.tsx", "DeBERTa-v3 in section 4.2", "ModernBERT-large with correct F1 figures"),
    ("DocumentationTab.tsx", "API URLs to localhost:5000/5001", "Live HF Space URLs"),
    ("ReportsTab.tsx", "61.07% shown as Task 1 Macro F1", "Replaced with honest model progression"),
    ("ModelDevelopment.tsx", "No leaderboard, no model history", "Full 7-row leaderboard V1 to Final"),
]
for r in content_rows:
    row = table2.add_row().cells
    for i, v in enumerate(r): row[i].text = v

para("")
para("New code snippets added to CodeTab.tsx:")
bullet("Company-disjoint split (GroupShuffleSplit by CompanyId) — the single most important methodological change")
bullet("ModernBERT-large training loop — Colab A100, fp16, epoch 3 to 70.29% Macro F1")
para("New content added to ModelDevelopment.tsx:")
bullet("7-row model leaderboard (V1 leaked 88.90% to Final locked 75.0%) with animated bars")
bullet("4-level cascade architecture diagram (Sector to Group to Industry to Sub-industry)")
bullet("Four key innovations section (company-disjoint split, GECS taxonomy anchoring, ModernBERT-large, calibration audit)")
bullet("Three result cards with all disclosed numbers (75.0% headline, 77.51% upper bound, 73.96% CV, ~76% structural ceiling)")
para("docs/model_version_history.md rewritten with the full 14-version honest progression, calibration audit breakdown, Task 2 progression, and cascade architecture diagram.")

# -------- 4. CHALLENGES --------
H2("4. Challenges encountered")

H3("4.1 HF Space routing confusion")
para(
    "The HF platform proxy intercepts some paths before they reach the application container. Only routes "
    "under /api/ that match Gradio's known endpoint structure are reliably forwarded. Bare /predict and custom "
    "aliases like /api/predict_legendary returned HF's own 404 HTML page regardless of being registered in "
    "the FastAPI app. This took multiple debugging cycles to isolate — the error message (DeBERTa server error "
    "(404): <!DOCTYPE html>...) was misleading because it contained HF's platform HTML rather than an "
    "application error. The fix was to standardize all proxy routes on /api/predict and verify the HF Space "
    "was set to public."
)

H3("4.2 framer-motion v12 SSR incompatibility")
para(
    "framer-motion v12 introduced internal changes to its CSS value-type registry that break server-side "
    "rendering for specific animation patterns: spring transitions with scale: 0 to 1, and backgroundPosition "
    "animations with multi-value strings. The Next.js App Router renders all client components on the server "
    "as initial HTML (SSR), so 'use client' alone does not prevent the crash — the ssr: false dynamic import "
    "pattern is required. The error message (TypeError: Cannot destructure property 'base' of 'f[c]') was "
    "opaque and took inspection of the framer-motion internals to trace to its root cause."
)

H3("4.3 The monitoring dashboard was displaying the leaked F1")
para(
    "The F1 stability graph in MonitoringTab.tsx had its oscillation range set to (86.0, 87.5) — the "
    "contaminated V1 result. This was displaying the discredited number on the production site. Caught during "
    "the content audit and corrected to (74.5, 75.5). This serves as a reminder that front-end content can "
    "carry stale claims just as silently as model evaluation scripts."
)

# -------- 5. PLAN --------
H2("5. Plan for final submission")
numbered("Draft the final capstone report incorporating both Task 1 (75.0%) and Task 2 (55.44%) results with the leakage audit, calibration audit, and structural ceiling analysis as the methodology backbone.")
numbered("Confirm HF Spaces remain public and warm; verify the full demo flow from frontend form submission through HF Space inference and back.")
numbered("Prepare submission package: final report, slide deck, demo URL, GitHub repo link.")
numbered("Write the project conclusion section: what generalizes beyond this dataset, what the honest methodology contributes to future cohorts.")

# -------- 6. REFLECTION --------
H2("6. Reflection")
para(
    "The dominant theme of Week 8 was closing gaps between the research result and the live artifact. Writing "
    "a model result is different from shipping it — deployment introduced a new class of failure (private "
    "space, wrong paths, SSR crashes, timeout mismatches) that had nothing to do with machine learning and "
    "everything to do with engineering discipline. The same rigor applied to model evaluation — verify the "
    "claim, trace the number, disclose the methodology — applied directly to the site content. Finding the F1 "
    "stability graph displaying 86% on the production dashboard was the clearest example: an honest result can "
    "be silently misrepresented by a single unchecked constant. The audit corrected it."
)
para(
    "Task 2 at 55.44% is below Task 1, and the reasons are structural: severe long-tail, "
    "hierarchy-propagated error, and a 428-class space that is mathematically harder than 145 classes with "
    "the same training volume. The number is honest, fully disclosed, and defensible — which is the standard "
    "established for Task 1 from the beginning."
)

hr()

meta("Time spent this week", "approximately 36 hours (Task 2 pipeline, HF Space deployment, Vercel deployment and debugging, frontend content audit, report writing)")
meta("Key artifacts produced this week", "hf_space/app.py (production), Vercel deployment (live), WEEK_8_REPORT.md, updated docs/model_version_history.md, all frontend tab and page content corrections")
meta("Status", "Task 1 locked 75.0% Macro F1 / 91.4% top-3. Task 2 locked 55.44% Macro F1 / 428 classes. Full production stack live. Final report drafting begins.")

out = r"C:\Users\akash\Desktop\capstone MGT 599\WEEK_8_REPORT.docx"
doc.save(out)
print(f"Saved: {out}")
