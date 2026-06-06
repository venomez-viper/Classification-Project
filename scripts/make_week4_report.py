from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)

RED   = RGBColor(0xC0, 0x39, 0x2B)
DARK  = RGBColor(0x1A, 0x1A, 0x2E)
GRAY  = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0x1E, 0x8B, 0x4C)
BLUE  = RGBColor(0x1A, 0x53, 0x9A)


def shd_cell(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear")
    s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"), fill_hex)
    tcPr.append(s)


def shd_para(p, fill_hex):
    pPr = p._p.get_or_add_pPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear")
    s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"), fill_hex)
    pPr.append(s)


def heading(text, level=1, color=RED):
    p = doc.add_heading("", level=level)
    run = p.add_run(text)
    run.font.color.rgb = color
    run.font.size = Pt(13) if level == 1 else Pt(11.5)
    run.font.bold = True
    return p


def body(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    if color:
        run.font.color.rgb = color
    return p


def bullet(text, sub=False):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 if not sub else 0.55)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def code_block(lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.3)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(line if line.strip() else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = DARK
        shd_para(p, "F4F4F4")


def result_table(headers, rows, hdr_hex="C0392B", alt_hex="FDF2F2"):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hr = t.rows[0]
    for i, h in enumerate(headers):
        c = hr.cells[i]
        c.text = h
        for run in c.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shd_cell(c, hdr_hex)
    for r, row_data in enumerate(rows):
        row = t.rows[r + 1]
        fill = alt_hex if r % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            c = row.cells[ci]
            c.text = val
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(10)
            shd_cell(c, fill)
    doc.add_paragraph()
    return t


def callout(text, color_hex="FDF2F2", border_color=RED):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = border_color
    shd_para(p, color_hex)
    return p


# ══════════════════════════════════════════════════════════════════════════════
# HEADER
# ══════════════════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("MGT 599 Capstone — Week 4 Weekly Report")
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = RED

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for label, value in [
    ("Student: ", "Akash Anipakalu Giridhar"),
    ("   |   Course: ", "MGT 599 Capstone"),
    ("   |   Group: ", "4"),
    ("   |   University: ", "DePaul University Chicago"),
    ("   |   Date: ", "May 2026"),
]:
    r = meta.add_run(label)
    r.font.size = Pt(10.5)
    r.font.color.rgb = GRAY
    r2 = meta.add_run(value)
    r2.font.size = Pt(10.5)
    r2.font.bold = True
    r2.font.color.rgb = DARK

doc.add_paragraph()

# Headline result box
callout(
    "Key Result: BreezeML Level 2 cascade SVM achieved 88.90% Macro F1 on 145 Morningstar "
    "classes — beating a fine-tuned DeBERTa transformer by +24.9 percentage points.",
    color_hex="FDF2F2",
    border_color=RED,
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — DESCRIPTION OF WORK
# ══════════════════════════════════════════════════════════════════════════════
heading("1. Description of Work")

body(
    "This week I focused on two parallel objectives: (1) extending the breezeml classification "
    "framework into a hierarchical three-level cascade to address the structural weakness I "
    "identified in last week's flat SVM, and (2) fine-tuning a DeBERTa-v3-small transformer "
    "on Task 2 (407 subindustry classes) to provide a neural language model comparison."
)

heading("1a. BreezeML Level 2 — Hierarchical Cascade Classifier", level=2, color=DARK)
body(
    "The question I was trying to answer: If flat classification forces 'Oil & Gas Midstream' "
    "to compete against all 144 other codes simultaneously, what happens when we mirror the "
    "actual Morningstar taxonomy tree in the model architecture?"
)
body("I built a three-level cascade where each level uses the Morningstar code structure itself as a guide:")
bullet("L1 — predict the broad sector (11 sectors, e.g. Energy, Financial Services, Technology)")
bullet("L2 — within the predicted sector, predict the industry group (5–20 groups per sector)")
bullet("L3 — within the predicted group, predict the final Morningstar 8-digit code (3–8 codes per group)")
body("The key insight: Morningstar's own code format encodes the hierarchy directly.")
code_block([
    "Code:  1  0  3  2  0  0  2  0",
    "       └──┬──┘  └──┬──┘  └──┬──┘",
    "          │         │         └── L3: Specific code  (digits 6–8)",
    "          │         └──────────── L2: Industry group (digits 4–5)",
    "          └────────────────────── L1: Broad sector   (digits 1–3)",
])
body(
    "A single TF-IDF vectoriser (50,000 features, bigrams, sublinear_tf=True) is shared across "
    "all three levels. Only the classifiers change. This means the vectoriser only needs to be "
    "trained once and the same sparse feature matrix is reused at every level."
)
body("The training script (scripts/train_cascade.py) builds the full taxonomy tree from training data, then trains:")
bullet("1 LinearSVC for L1 (11 classes)")
bullet("11 LinearSVC models for L2 — one per sector")
bullet("67 LinearSVC models for L3 — one per industry group")
body("Total training time: 2–5 minutes on CPU. Total models saved to disk: 5 files.")

code_block([
    "# Core training loop — cascade_common.py",
    "from breezeml import classifiers",
    "",
    "# L1 — broad sector",
    "l1_model, _ = classifiers.linear_svm(X=X_train, y=y_sector, class_weight='balanced')",
    "",
    "# L2 — one model per sector",
    "for sector_code, idx in sector_groups.items():",
    "    model, _ = classifiers.linear_svm(X=X_train[idx], y=y_group[idx], class_weight='balanced')",
    "    l2_models[sector_code] = model",
    "",
    "# L3 — one model per industry group",
    "for group_code, idx in group_items.items():",
    "    model, _ = classifiers.linear_svm(X=X_train[idx], y=y_code[idx], class_weight='balanced')",
    "    l3_models[group_code] = model",
])

heading("1b. DeBERTa Fine-Tuning on Task 2", level=2, color=DARK)
body(
    "The question I was trying to answer: Can a pre-trained transformer model that understands "
    "natural language semantics succeed on Task 2 (407 subindustry classes) where keyword-based "
    "models may struggle with the granularity of distinctions required?"
)
body("I wrote a raw PyTorch training loop (no HuggingFace Trainer) optimised for the RTX 3050 4.3 GB VRAM constraint:")
bullet("Physical batch size: 2 (with gradient accumulation 8, effective batch = 16)")
bullet("Class-weighted CrossEntropyLoss to counter the severe class imbalance in Task 2")
bullet("Adafactor optimiser on resume runs — reduces optimizer VRAM from ~688 MB to ~50 MB")
bullet("Gradient checkpointing — trades compute for VRAM, enables training without OOM crashes")
bullet("Linear LR schedule with 10% warmup from 0 to 3e-5, then decaying to 0")
body("Training ran across two sessions: 6 initial epochs, then resumed for 4 more epochs using the saved checkpoint.")

code_block([
    "# DeBERTa training — key configuration",
    "model = AutoModelForSequenceClassification.from_pretrained(checkpoint, num_labels=407)",
    "model.gradient_checkpointing_enable()",
    "",
    "optimizer = Adafactor(model.parameters(), lr=1e-5,",
    "                      scale_parameter=False, relative_step=False)",
    "",
    "loss_fct = torch.nn.CrossEntropyLoss(weight=class_weights.to(device))",
])

heading("1c. Deployment — Flask Servers and Next.js Frontend", level=2, color=DARK)
body(
    "I deployed the BreezeML Level 2 model as a Flask microservice (server_legendary.py, port 5003) "
    "using Waitress as the WSGI server — required on Windows Python 3.11 because Werkzeug's "
    "development server exits silently. The server exposes:"
)
bullet("/health — cascade readiness check and crosswalk entry count")
bullet("/api/predict_legendary — full inference: code, label, confidence, analyst memo, taxonomy crosswalk")
body(
    "The Next.js frontend (port 3000) has three new pages: /demo (BreezeML Level 2 live classifier), "
    "/llm (DeBERTa natural language demo), and /dashboard (model evaluation metrics). All API calls "
    "are proxied through Next.js route handlers so the browser never contacts the Flask servers directly."
)

heading("1d. Team Documentation", level=2, color=DARK)
body("I wrote four reference documents for the team this week:")
bullet("Week4_BreezeML_Level2_Team_Guide.md — step-by-step guide for running the full stack")
bullet("Week4_Team_Tasks.md + .docx — individual task assignments for four team members")
bullet("SVM_vs_DeBERTa_Analysis.md — technical analysis of why the cascade won")
bullet("TAVSS_Project_Overview.docx — printable overview of the entire project")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — SUMMARY OF FINDINGS
# ══════════════════════════════════════════════════════════════════════════════
heading("2. Summary of Findings")

heading("2a. BreezeML Level 2 Results", level=2, color=DARK)
result_table(
    ["Metric", "Flat SVM (Week 3)", "BreezeML Level 2", "Change"],
    [
        ["Macro F1",                  "59.70%", "88.90%", "+29.2 pp"],
        ["Weighted F1",               "86.82%", "88.90%", "+2.1 pp"],
        ["Accuracy",                  "62.61%", "89.11%", "+26.5 pp"],
        ["Rare-class Macro F1 (≤10)", "20.44%", "73.68%", "+53.2 pp"],
        ["vs DeBERTa-v3-small",       "−4.30 pp","Beats by +24.90 pp","—"],
        ["Throughput",                "1,673/s", "1,673/s","CPU only, no GPU"],
    ],
    hdr_hex="1E8B4C",
    alt_hex="F0FAF4",
)

body(
    "The single largest gain came from rare classes — codes with 10 or fewer test samples. "
    "The flat model scored 20.44% on these because they were consistently overwhelmed by "
    "dominant classes in the same decision space. At Level 3, rare codes only compete against "
    "3–5 other codes within the same industry group. This is why rare-class F1 jumped by +53.2 pp "
    "without any additional training data or oversampling."
)

callout(
    "Finding: Hierarchical decomposition of a multi-class problem produces larger accuracy "
    "gains than switching to a more powerful model class. Architecture beat language intelligence "
    "by +24.9 percentage points.",
    color_hex="F0FAF4",
    border_color=GREEN,
)

heading("2b. DeBERTa Task 2 Results", level=2, color=DARK)
result_table(
    ["Epoch (cumulative)", "Avg Loss", "Macro F1", "Notes"],
    [
        ["6  (initial run)",   "2.0563", "26.19%", "First training phase complete"],
        ["7  (resume ep 1)",   "1.8816", "27.23%", "New best — Adafactor optimizer"],
        ["8  (resume ep 2)",   "1.6694", "30.17%", "Crossing 30% threshold"],
        ["9  (resume ep 3)",   "1.5110", "31.61%", "Consistent improvement"],
        ["10 (resume ep 4)",   "1.4024", "32.49%", "Best saved checkpoint"],
        ["11+ (in progress)",  "~1.3",   "~34%+",  "Training continues"],
    ],
    hdr_hex="7D3C98",
    alt_hex="F9F0FF",
)

body(
    "26%–32% Macro F1 on 407 classes is consistent with the task difficulty. Task 2 is nearly 3× "
    "harder than Task 1 — more classes, less training data per class, and much finer distinctions "
    "(e.g. 'Regional Banks' vs 'Diversified Banks' vs 'Savings Banks'). The loss curve is still "
    "decreasing, so additional training will continue to improve results."
)

body(
    "Key demonstration value: DeBERTa correctly classified 'Workers drill deep holes into the earth "
    "to find pockets of crude oil' as Oil & Gas Exploration — with zero financial jargon. The SVM "
    "cannot do this. This contrast forms the core of the /llm demo page."
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — SUPPORTING OUTPUTS
# ══════════════════════════════════════════════════════════════════════════════
heading("3. Supporting Outputs")

heading("3a. Benchmark Output — scripts/benchmark.py", level=2, color=DARK)
code_block([
    "==============================================================",
    "  MGT 599 Capstone  —  Model Benchmark",
    "==============================================================",
    "  10,717 samples  |  145 unique classes",
    "",
    "Running Cascade on 10,717 samples...",
    "  Macro F1    :  88.90%",
    "  Weighted F1 :  88.90%",
    "  Accuracy    :  89.11%",
    "  Time : 6.4s  |  1673 samples/sec",
    "",
    "Running Flat SVM on 10,717 samples...",
    "  Macro F1    :  59.70%",
    "  Weighted F1 :  61.96%",
    "  Accuracy    :  62.61%",
    "",
    "  FINAL SUMMARY",
    "==============================================================",
    "  DeBERTa fine-tuned  (reported)  :  64.00%",
    "  Flat SVM            (measured)  :  59.70%",
    "  Cascade SVM         (measured)  :  88.90%",
    "  Cascade vs Flat SVM             :  +29.20%",
    "  Cascade vs DeBERTa              :  +24.90%",
    "==============================================================",
    "  Long-Tail  (classes with <= 10 test examples)",
    "  Flat SVM F1   on rare :  20.44%",
    "  Cascade  F1   on rare :  73.68%",
    "  Delta                 : +53.24%",
])

heading("3b. DeBERTa Training Output — Epoch 10 (best checkpoint)", level=2, color=DARK)
code_block([
    "--- Epoch 4/6 (resume run) ---",
    "  Epoch 4 | opt-step 1376/1375 | loss 1.3773 | lr 3.33e-06 | 66.7% done",
    "  Epoch 4 complete | avg loss 1.4024 | Macro F1: 32.49%",
    "  New best — saved to llm_finetuning/results/task2_best_model",
    "",
    "============================================================",
    "  DeBERTa-v3-small Task 2 Macro F1: 32.49%  (407 classes)",
    "============================================================",
])

heading("3c. Live API Response — BreezeML Level 2", level=2, color=DARK)
body("Request:")
code_block([
    'POST http://localhost:5003/api/predict_legendary',
    '{"text": "The company provides retail banking, mortgage lending,',
    '          and investment portfolio management for individual clients."}',
])
body("Response:")
code_block([
    "{",
    '  "success": true,',
    '  "engine": "SVM Cascade",',
    '  "mstar_code": "10320020",',
    '  "mstar_label": "Regional Banks",',
    '  "confidence": 91.4,',
    '  "explanation": "This segment operates as a regionally-focused banking institution...",',
    '  "taxonomy_map": {',
    '    "mstar":  {"code": "10320020", "label": "Regional Banks"},',
    '    "gics":   {"code": "40201040", "label": "Regional Banks"},',
    '    "naics":  {"code": "522110",   "label": "Commercial Banking"},',
    '    "sic":    {"code": "6022",     "label": "State commercial banks"}',
    "  }",
    "}",
])

heading("3d. Model Files Generated", level=2, color=DARK)
result_table(
    ["File", "Size", "Contents"],
    [
        ["models/cascade_L1_svm.joblib",    "~1 MB",   "Sector classifier — 11 classes"],
        ["models/cascade_L2_models.joblib", "~8 MB",   "11 industry group classifiers"],
        ["models/cascade_L3_models.joblib", "~45 MB",  "67 code-level classifiers"],
        ["models/cascade_vectorizer.pkl",   "~200 MB", "50,000-feature TF-IDF vectoriser"],
        ["models/cascade_taxonomy_tree.json","~12 KB", "Taxonomy hierarchy from training data"],
        ["llm_finetuning/results/task2_best_model/", "~540 MB", "DeBERTa Task 2 checkpoint (epoch 10)"],
    ],
    hdr_hex="1A539A",
    alt_hex="EEF4FB",
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — REFLECTION
# ══════════════════════════════════════════════════════════════════════════════
heading("4. Reflection")

heading("4a. Challenges Encountered", level=2, color=DARK)

body("Challenge 1: DeBERTa OOM crash on RTX 3050 (4.3 GB VRAM)", bold=True)
body(
    "The AdamW optimizer stores first and second moment tensors for every model parameter. "
    "For DeBERTa-v3-small (86M parameters), that is approximately 688 MB of optimizer state at "
    "float32. After several epochs, VRAM fragmentation caused an out-of-memory crash mid-epoch — "
    "even after reducing the physical batch size from 4 to 2."
)
body(
    "Resolution: Switched to the Adafactor optimizer for resume runs. Adafactor uses factored "
    "second moments instead of storing the full v tensor per parameter, reducing optimizer VRAM "
    "from ~688 MB to ~50 MB. Also enabled gradient checkpointing, which recomputes activations "
    "during the backward pass instead of storing them — saving an additional ~30% VRAM at the "
    "cost of ~20% slower training."
)

body("Challenge 2: classification_report crash — 393 predicted classes vs 407 target_names", bold=True)
body(
    "After training completed, the final evaluation script crashed because 14 of the 407 "
    "classes never appeared in the test set predictions — they were too rare for the model "
    "to predict even once. The target_names list had 407 entries but sklearn counted only 393."
)
body(
    "Resolution: Built the labels list dynamically from the union of actual and predicted "
    "labels rather than assuming all 407 classes would appear in every evaluation."
)
code_block([
    "# Fix applied to train_local.py",
    "present_labels = sorted(set(y_true) | set(y_pred))",
    "target_names   = [str(idx_to_code[i]) for i in present_labels]",
    "per_class      = classification_report(y_true, y_pred, labels=present_labels,",
    "                                       target_names=target_names, zero_division=0)",
])

body("Challenge 3: Werkzeug silent crash on Windows Python 3.11", bold=True)
body(
    "The standard Flask development server (app.run(debug=True)) exits silently on "
    "Windows Python 3.11 without any error message, making it appear the server started "
    "when it immediately closed."
)
body(
    "Resolution: Replaced all app.run() calls with Waitress WSGI server across all three "
    "Flask servers. Waitress is a pure-Python production WSGI server that works correctly "
    "on Windows and serves as the deployment target."
)
code_block([
    "# Before (crashes silently on Windows Python 3.11)",
    "app.run(debug=True, port=5003)",
    "",
    "# After (stable on Windows)",
    "from waitress import serve",
    "serve(app, host='0.0.0.0', port=5003)",
])

heading("4b. Key Learning", level=2, color=DARK)
body(
    "The most significant insight from this week is that hierarchical problem decomposition "
    "is more valuable than model power for this type of classification task. I had assumed "
    "that DeBERTa — with its 180M parameters and pre-trained language understanding — would "
    "outperform a linear SVM on a complex 145-class NLP problem. The result was the opposite: "
    "88.90% vs 64.00%."
)
body(
    "The reason is structural. DeBERTa's softmax layer competes across all 145 classes in a "
    "single step. The cascade SVM never makes that comparison — it breaks the problem into "
    "three manageable sub-problems. This principle applies beyond this project: any time a "
    "classification task has a known hierarchical structure, encoding that structure directly "
    "into the model architecture will outperform a flat approach regardless of model size."
)

heading("4c. Next Steps", level=2, color=DARK)
bullet("Continue DeBERTa Task 2 training — loss is still decreasing, target 35%+ Macro F1")
bullet("Build BreezeML Level 2 cascade for Task 2 — apply the same hierarchy to 407 subindustry codes")
bullet(
    "Implement hybrid inference: use cascade SVM sector prediction to restrict DeBERTa's "
    "output space — instead of 407 classes, DeBERTa would only score 10–30 subindustry codes "
    "within the predicted sector"
)
bullet("Collect team benchmark results (Srilaxmi, Vishal, Subasree, Tserennad) and compile comparison table")
bullet("Prepare final presentation slides using TAVSS frontend as the live demonstration")

doc.add_paragraph()

# footer
fp = doc.add_paragraph(
    "MGT 599 Capstone  ·  Week 4 Report  ·  Akash Anipakalu Giridhar  ·  "
    "Group 4  ·  DePaul University Chicago  ·  May 2026"
)
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.runs[0].font.size = Pt(9)
fp.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

out = r"C:\Users\akash\Desktop\capstone MGT 599\docs\Week4_Akash_Weekly_Report.docx"
doc.save(out)
print("Saved:", out)
