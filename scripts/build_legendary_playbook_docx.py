"""Render the Legendary Playbook as a polished Word doc — internal strategy, not for submission."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
OUT  = ROOT / "docs" / "LEGENDARY_PLAYBOOK.docx"

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
TEAL = RGBColor(0x0E, 0x6B, 0x6E)
GOLD = RGBColor(0xD4, 0xA9, 0x3F)
GRAY = RGBColor(0x3D, 0x49, 0x51)

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.2); s.bottom_margin = Cm(2.2)
    s.left_margin = Cm(2.0); s.right_margin = Cm(2.0)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def heading(text, level=2, color=NAVY):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = color
        if level == 1: r.font.size = Pt(20)
        elif level == 2: r.font.size = Pt(15)
        else: r.font.size = Pt(12)


def para(text, bold=False, italic=False, size=11, color=GRAY):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    r.font.color.rgb = color
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text); r.font.size = Pt(11); r.font.color.rgb = GRAY


def quote_block(text, color=TEAL):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(20)
    p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "8"); left.set(qn("w:color"), "0E6B6E")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    r.italic = True; r.font.size = Pt(11.5); r.font.color.rgb = color


def shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def styled_table(rows, header_hex="1F3A5F", col_widths_in=None):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = t.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10); run.font.color.rgb = GRAY
            if r_idx == 0:
                run.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                shade_cell(cell, header_hex)
            elif r_idx % 2 == 1:
                shade_cell(cell, "F7F7F7")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths_in:
        for c_idx, w in enumerate(col_widths_in):
            for row in t.rows:
                row.cells[c_idx].width = Pt(w * 72)
    doc.add_paragraph()


def hrule():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom"); bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6"); bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom); pPr.append(pBdr)


# ─── Cover ───
p = doc.add_paragraph()
r = p.add_run("INTERNAL STRATEGY  ·  NOT FOR SUBMISSION")
r.font.size = Pt(9); r.bold = True; r.font.color.rgb = GOLD
doc.add_paragraph()

heading("The Legendary Playbook for Monday", level=1)
para("MGT 599 Capstone · Final Presentation Strategy", italic=True)
para("Group 4 · Lead: Akash Anipakalu Giridhar")
para("Audience: Morningstar Reference Entity Data (RED) team")
para("Date prepared: May 11, 2026  ·  For submission: May 18, 2026")
doc.add_paragraph()
quote_block("Use this to drive every decision in the final week. F1 alone will not impress a senior data science team. Judgment, production thinking, and domain humility will.")
hrule()

heading("The Goal", level=2)
para("Walk out of the Monday presentation having made the Morningstar people feel like they should hire the lead. F1 number alone doesn't do that. Senior judgment, production thinking, and domain humility do.")
para("This document is the playbook for getting there.")
hrule()

# ─── 1 ───
heading("1. The Thesis — one sentence Morningstar will remember", level=2)
quote_block('"Industry classification is not 100% automatable. The right product is an analyst-first system where the model handles the obvious cases, defers on hard ones, and explains every decision using the company\'s own taxonomy."')
para("This is your positioning. Honest, enterprise-mature, and the opposite of what 90% of capstone teams will say. They'll sell 'we built AI.' You sell 'we built a tool your analysts will actually use.'")
para("Say this sentence twice — once at minute 1, once at minute 30.", bold=True)
hrule()

# ─── 2 ───
heading("2. The Narrative Arc — three acts, not 'here's what we did'", level=2)
heading("Act I — The Discovery (Weeks 1–3)", level=3)
quote_block('"We built a TF-IDF cascade and got 88.90% Macro F1. We almost shipped it."')
heading("Act II — The Audit (Week 4)", level=3)
quote_block('"Then we audited our own evaluation pipeline. We found that 97.2% of our test set was inside our training set. The 88.90% was leaked memorization. The honest number was 60%. We rebuilt everything from scratch."')
heading("Act III — The System (Weeks 5–7)", level=3)
quote_block('"With honest evaluation in place, we built a production-ready classification system grounded in your own 2019 GECS taxonomy document. Here\'s the live demo."')
para("This arc is why they hire you.", bold=True)
para("Every junior data scientist builds models. The senior ones audit their own work and tell unflattering truths. Lead with the audit, not the F1.")
hrule()

# ─── 3 ───
heading("3. What You Are Actually Selling — a product, not a model", level=2)
para("A model is a .joblib file. A product is:")
bullet("An API analysts can call from their existing workflow")
bullet("A UI that explains its reasoning in plain English using the GECS definitions Morningstar wrote")
bullet("A confidence number that means something (calibrated, not softmax-on-margin)")
bullet("An audit trail of every prediction")
bullet("A path to retrain when the taxonomy updates")
bullet("A cost-per-prediction number you can defend")
bullet("A live URL anyone in the room can hit on their phone")
para("If your final 30 seconds is 'here's how it integrates with RED's workflow tomorrow morning,' you have stopped being a student to them.", bold=True)
hrule()

# ─── 4 ───
heading("4. The Five Legendary Moves We Make This Week", level=2)

heading("Move 1 — Frame the audit as your first contribution, not a footnote", level=3)
para("The leakage discovery is your single strongest signal of senior judgment. Every other team will downplay theirs (or will not have caught one). Lead with yours. Open the presentation with the 88.90% → 60% slide. Make Morningstar sit up.")

heading("Move 2 — Ground every prediction in the regulator's own words", level=3)
para("Nobody else will think to parse the Morningstar 2019 GECS PDF and use its 145 official definitions as semantic anchors. We have. When the demo predicts a code, it cites the exact phrase from the official definition that matched.")
quote_block('This is domain humility. It says: "We didn\'t invent labels. We used yours."')

heading("Move 3 — Build the analyst-override workflow, not full automation", level=3)
para("Most teams pretend their model is 'ready for production.' Pre-empt the obvious objection — 'What about the hard cases?' — by designing for it.")
para("Your system shows top-3 candidates with calibrated confidence and an 'Override' button. The message: the analyst is the final authority; the model makes them 5× faster. That is the right product story for RED.")

heading("Move 4 — Deploy it as a live URL on Hugging Face Spaces", level=3)
para("Walk into the room with a URL. Have the Morningstar rep pull out their phone and type a description. The model responds in 200ms with reasoning.")
para("You stop being a student the moment they can interact with your system from their own device. That is the moment.", bold=True)

heading("Move 5 — Name what you cannot solve, and recommend the workflow", level=3)
para("The hardest class, 31030010 Diversified Industrials, will not hit 85% F1. Even humans disagree on conglomerate boundaries.")
para("Most teams will hide this. You surface it explicitly:")
quote_block('"Class 31030010 is structurally hard. Even humans disagree on conglomerate boundaries. We recommend routing predictions for this class through RED\'s senior-analyst review queue. Our system flags them automatically based on segment-count and revenue-dispersion features."')
para("That slide says: I understand your business. I understand what ML can and cannot do. I built the workflow around the limitation. Hireable behavior.", bold=True)
hrule()

# ─── 5 ───
heading("5. The Demo Flow — five minutes that wins the room", level=2)
styled_table([
    ["t (min)", "What happens"],
    ["0:00", "Open the live Hugging Face Space URL on the projector."],
    ["0:15", "Paste a real description: 'Operates regional retail banks in the Midwest with ~50 branches focused on commercial lending.'"],
    ["0:45", "Show the response: Task 1 + Task 2 + official GECS definition quoted + top-3 alternatives + processing trace + ~87% calibrated confidence."],
    ["1:30", "Click 'Show reasoning' → display the chain: 'Segment text mentions retail banking + regional + commercial lending → Sector 103 Financial Services → Group 10320 Banks → Code 10320020 Banks—Regional.'"],
    ["2:30", "Paste an ambiguous conglomerate description on purpose. Show the system DEFERRING with low confidence: 'Confidence below threshold (54%). Recommend analyst review.' This moment shows judgment, not performance."],
    ["3:30", "Switch to the /metrics page. Latency p95, predictions logged, confidence histogram. 'This is production-ready, not a notebook.'"],
    ["4:00", "Ask the Morningstar rep to type their own example. Let them play."],
    ["4:30", "Close with the closing line (Section 7)."],
])
hrule()

# ─── 6 ───
heading("6. What We Do NOT Do This Week", level=2)
styled_table([
    ["Don't", "Why"],
    ["Oversell the F1 number",
     "If the number is 72%, say 72%. Morningstar will see through inflation. Your audit story makes a real 72% more credible than a fake 85%."],
    ["Run more experiments after Tuesday",
     "Lock the model Tuesday night. Spend Wednesday–Friday on packaging, not science. Every team underestimates this."],
    ["Use buzzwords without code behind them",
     "If you say 'RAG,' show it in the architecture diagram. If you say 'calibrated probabilities,' point to CalibratedClassifierCV in the code. Empty buzzwords get caught in Q&A."],
    ["Hide failure cases",
     "Walk through three predictions: one easy, one hard but right, one where the system defers. The deferral wins more credibility than the success."],
    ["Pitch alone",
     "If the team's there, give one slide each. Morningstar evaluates leadership and team coordination."],
])
hrule()

# ─── 7 ───
heading("7. The Closing Line — say this verbatim", level=2)
para("After the demo, after the questions, look the Morningstar rep in the eye and say:")
quote_block('"We didn\'t try to replace your analysts. We tried to build the tool we\'d want as one of them. Everything in this system is grounded in your taxonomy, calibrated honestly, and deployable on infrastructure you already have. We\'d love to hear what would need to change for this to land in RED\'s workflow."')
para("That last sentence — asking for feedback on production fit — flips the energy in the room. You stop being graded. You become a candidate having a conversation with a hiring manager.", bold=True)
hrule()

# ─── 8 ───
heading("8. What the Lead (Akash) Personally Does Before Monday", level=2)
bullet("Re-read the GECS PDF cover to cover. When you cite it in the presentation, cite the exact page. Morningstar people will hear it.")
bullet("Pre-rehearse with two real Morningstar coverage companies. Pick two public companies you know. Practice classifying them out loud, including the reasoning. Your demo will be 5× sharper.")
bullet("Practice the leakage-audit slide alone in the mirror. Get the timing tight: 'We almost shipped 88.90%. Then we caught a 30-point leak. Here's what we learned about audit discipline.' 90 seconds, no notes.")
bullet("Sleep before Monday. A clear head delivering a 72% honest number beats a tired one delivering 85%.")
hrule()

# ─── 9 ───
heading("9. What This Buys You by Monday", level=2)
para("When you walk in, you have:")
bullet("A live URL anyone in the room can touch")
bullet("A 90-second audit story that signals senior judgment")
bullet("Reasoning traces backed by Morningstar's own text")
bullet("Honest performance numbers with no asterisks")
bullet("A workflow design that respects analyst authority")
bullet("A closing line that invites a hiring conversation")
para("That is what gets remembered. Not the F1.", bold=True)
hrule()

# ─── 10 ───
heading("10. Anchor Truths to Stay Calm Under Pressure", level=2)
para("Whatever happens this week, these stay true:")
bullet("We caught a 30-point leakage in our own work. Most teams won't.")
bullet("We built methodology rigor a junior wouldn't have shown.")
bullet("We grounded every prediction in the regulator's own document.")
bullet("We delivered an analyst-friendly product, not a science experiment.")
bullet("The number we deliver is real. The work to get there is documented end-to-end.")
para("If F1 hits 75% on Monday — great. If it hits 72% — also great. The story is the same. The decision-quality is what they're hiring for.", bold=True)

doc.add_paragraph()
foot = doc.add_paragraph(); foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot.add_run("Prepared by Akash Anipakalu Giridhar  ·  MGT 599 Capstone  ·  DePaul University Chicago")
r.font.size = Pt(9); r.italic = True; r.font.color.rgb = GRAY
foot2 = doc.add_paragraph(); foot2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = foot2.add_run("Last updated: May 11, 2026  ·  INTERNAL STRATEGY — NOT FOR SUBMISSION")
r2.font.size = Pt(8); r2.italic = True; r2.font.color.rgb = GOLD

doc.save(OUT)
print(f"Wrote {OUT}")
