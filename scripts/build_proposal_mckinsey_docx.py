"""Build the McKinsey-style proposal Word doc with embedded exhibits and code."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
EXH  = ROOT / "docs" / "proposal_exhibits"
OUT  = ROOT / "docs" / "Initial_Proposal_McKinsey.docx"
OUT_TAVSS = ROOT / "docs" / "Initial_Proposal_TAVSS.docx"

# ── Palette ──
NAVY  = RGBColor(0x1F, 0x3A, 0x5F)
TEAL  = RGBColor(0x0E, 0x6B, 0x6E)
CORAL = RGBColor(0xE0, 0x78, 0x56)
GOLD  = RGBColor(0xD4, 0xA9, 0x3F)
SAGE  = RGBColor(0x7A, 0x9E, 0x7E)
GRAY  = RGBColor(0x3D, 0x49, 0x51)
LIGHT = "E8ECEF"
DARK_HEX = "1F3A5F"
TEAL_HEX = "0E6B6E"
CODE_HEX = "F4F4F0"

doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)


# ────────────────────────── Helper functions ──────────────────────────
def shade(p, hex_color):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def cell_border(cell, edges=("top","bottom","left","right"), size="4", color="BFBFBF"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in edges:
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def action_title(text):
    """McKinsey action title — declarative answer at the top of each section."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(2)
    # Top accent line
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single"); top.set(qn("w:sz"), "12")
    top.set(qn("w:space"), "6"); top.set(qn("w:color"), DARK_HEX)
    pBdr.append(top)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.size = Pt(15)
    run.font.color.rgb = NAVY
    run.bold = True


def subtitle(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = TEAL
    run.italic = True


def body(text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.bold = bold
    run.font.color.rgb = GRAY
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = GRAY


def section_marker(text):
    """Small uppercase section label."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run(text.upper())
    run.font.size = Pt(9)
    run.bold = True
    run.font.color.rgb = TEAL
    run.font.name = "Calibri"


def exhibit_caption(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY


def insert_image(filename, width_in=6.5, caption=None):
    img_path = EXH / filename
    if not img_path.exists():
        body(f"[image missing: {filename}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(width_in))
    if caption:
        exhibit_caption(caption)


def code_block(code: str, lang_label=""):
    if lang_label:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(lang_label)
        r.italic = True; r.font.size = Pt(8.5)
        r.font.color.rgb = TEAL
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(8)
    p.paragraph_format.space_after = Pt(10)
    shade(p, CODE_HEX)
    # Left border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "6"); left.set(qn("w:color"), TEAL_HEX)
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(8.6)
    run.font.color.rgb = GRAY


def kpi_tiles(tiles):
    """Row of KPI tiles. tiles = list of (value, label, color_hex)."""
    table = doc.add_table(rows=1, cols=len(tiles))
    table.autofit = False
    for i, (value, label, hex_fill) in enumerate(tiles):
        cell = table.rows[0].cells[i]
        cell.text = ""
        shade_cell(cell, hex_fill)
        cell_border(cell, color="FFFFFF", size="6")
        # Big value
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(value)
        r.font.size = Pt(20)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Label below
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(label)
        r2.font.size = Pt(8.5)
        r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r2.bold = True
    doc.add_paragraph()


def styled_table(rows, header=True, col_widths_in=None, header_hex="1F3A5F"):
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.color.rgb = GRAY
            if r_idx == 0 and header:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                shade_cell(cell, header_hex)
            elif r_idx % 2 == 1 and header:
                shade_cell(cell, "F7F7F7")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths_in:
        for c_idx, w in enumerate(col_widths_in):
            for row in table.rows:
                row.cells[c_idx].width = Inches(w)
    doc.add_paragraph()
    return table


# ─────────────────────────────────────────────────────────────────────────────
# COVER PAGE
# ─────────────────────────────────────────────────────────────────────────────
# Top label
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("MGT 599 CAPSTONE PROJECT  ·  Q2 2026")
r.font.size = Pt(9.5); r.bold = True
r.font.color.rgb = TEAL

doc.add_paragraph()
doc.add_paragraph()

# Title
p = doc.add_paragraph()
r = p.add_run("Initial Project Proposal")
r.font.size = Pt(13)
r.font.color.rgb = GRAY

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run("GECS Industry and Business Activity")
r.font.size = Pt(34)
r.font.color.rgb = NAVY
r.bold = True

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(20)
r = p.add_run("Classification using Machine Learning")
r.font.size = Pt(34)
r.font.color.rgb = NAVY
r.bold = True

# Subtitle / promise
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
r = p.add_run("An auditable, hierarchy-aware NLP pipeline for Morningstar's "
              "Reference Entity Data team — built on honest evaluation, "
              "grounded in the official GECS taxonomy.")
r.font.size = Pt(12.5)
r.italic = True
r.font.color.rgb = TEAL

# 3-month engagement note
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(40)
r = p.add_run("Three-month engagement: March 5 – June 6, 2026  ·  "
              "Source data: case-issued task1_gecs_classification_final.csv "
              "and task2_subindustry_classification_final.csv")
r.font.size = Pt(10)
r.italic = True
r.font.color.rgb = GRAY

# KPI tiles at the bottom of the cover
kpi_tiles([
    ("69.09%", "Current honest Macro F1\nV10 calibrated stack (Week 4)",  "1F3A5F"),
    ("88.90%", "Reported leaked Macro F1\nInvalid for decision-making",   "E07856"),
    ("75–80%", "Target operating range\nWeek 5–7 sprint goal",            "0E6B6E"),
    ("97.2%",  "Leaked test overlap\nCaught in our own audit",            "D4A93F"),
])

# Footer block
doc.add_paragraph()
p = doc.add_paragraph()
for line, bold in [
    ("Group 4 · DePaul University Chicago", True),
    ("Lead: Akash Anipakalu Giridhar", False),
    ("Team: Srilaxmi, Vishal, Subasree, Tserennad", False),
    ("Industry Partner: Morningstar — Reference Entity Data (RED) Team", False),
    ("May 10, 2026", False),
]:
    r = p.add_run(line + "\n")
    r.font.size = Pt(10.5)
    r.bold = bold
    r.font.color.rgb = GRAY

doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 0 — How to read
# ─────────────────────────────────────────────────────────────────────────────
section_marker("Section 0  ·  How to read this report")
action_title("The report is structured around the decision problem, the audit, and the path to a defensible model.")
subtitle("Read the action titles as the storyline. Each section is written to answer the executive question: "
         "what changed, why it changed, and what we are recommending next.")
body("This proposal follows a pyramid structure: a top-line recommendation, the analytical findings that support "
     "it, the audit that re-grounded our thinking, and the architecture we will build into the remaining sprint. "
     "Exhibits and code excerpts are interleaved with the narrative so that any single page can be lifted into a "
     "leadership review.")

styled_table([
    ["Section", "Action title"],
    ["1", "Prioritize honest hierarchy-aware modeling over chasing the leaked 88.90% result."],
    ["2", "GECS is a fine-grained taxonomy problem where small input errors compound downstream."],
    ["3", "The Week 4 audit corrected a 30-point performance illusion and reset the modeling path."],
    ["4", "The current best honest model reaches 69.09% Macro F1; results cluster around a structural ceiling."],
    ["5", "LongProfile contamination is the most likely root cause of the remaining performance ceiling."],
    ["6", "The next architecture should combine clean segment inputs with hierarchy-aware learning."],
    ["7", "Task 2 should exploit the deterministic Task 1 → Task 2 relationship, not act as a 450-class flat model."],
    ["8", "The remaining sprint should optimize for credibility, rare-class lift, and demo realism."],
    ["9", "The final submission should be positioned as an enterprise-grade ML audit and classification system."],
], col_widths_in=[0.7, 6.0])


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 1 — Recommendation
# ─────────────────────────────────────────────────────────────────────────────
section_marker("Section 1  ·  Recommendation")
action_title("Prioritize honest hierarchy-aware modeling over chasing the leaked 88.90% result.")
subtitle("The answer first: the credible baseline is 69.09% Macro F1, and the remaining sprint should focus "
         "on clean segment-only inputs, hierarchy-aware learning, and rare-class treatment.")

body("We should not present the 88.90% cascade result as evidence of model readiness. That number was "
     "produced by training on data that overlapped 97.2% with the test set; on the 305 truly-unseen rows the "
     "same model scored 81.73%, and on a clean re-evaluation it scored ~60%. Presenting the leaked number to "
     "Morningstar would invite a single follow-up question that would invalidate our work. We caught it first.")

body("We have therefore moved the project from a simple model-building exercise into a credible enterprise ML "
     "engagement: we define the data contract, audit the evaluation pipeline, build hierarchy-aware models on "
     "clean inputs, and communicate calibrated confidence.")

body("Three decisions follow from the audit:", bold=True)
bullet("Use strict train/test separation and preserve audit evidence for every headline metric.")
bullet("Reduce input contamination by testing segment-only text as the primary decision signal.")
bullet("Move from flat 145-class prediction toward hierarchy-aware modeling that predicts sector, group, and industry jointly.")

insert_image("exhibit_1_performance_journey.png", width_in=6.5,
              caption="Exhibit 1: The performance story changed materially after we removed training/test leakage.")


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 2 — Problem
# ─────────────────────────────────────────────────────────────────────────────
section_marker("Section 2  ·  Problem framing")
action_title("Morningstar GECS is a fine-grained taxonomy problem where small input errors compound downstream.")
subtitle("The case asks for scalable NLP classification across a deep hierarchy, not merely high accuracy on a "
         "convenient random split.")

body("The Reference Entity Data team maintains GECS as a four-level hierarchy used for peer comparison, "
     "portfolio construction, risk attribution, and reporting across Morningstar and PitchBook. "
     "Misclassification at the leaf level distorts every product downstream.")

insert_image("exhibit_2_hierarchy.png", width_in=6.5,
              caption="Exhibit 2: GECS rolls company descriptions into increasingly fine-grained industry and business-activity labels.")

body("The problem is challenging for three structural reasons:")
bullet("Class imbalance is material. The top-10 codes account for roughly 30% of records; the bottom 50 codes have fewer than 200 records each. A model that maximizes accuracy will under-serve rare codes.")
bullet("Multi-segment companies are common. About 35% of CompanyId values have segments mapping to different MstarGlobal codes. These are diversified conglomerates whose company-level descriptions span multiple sectors.")
bullet("Cascade error propagation. Errors at the sector level flow into all downstream levels. A naïve top-down cascade is structurally fragile.")

styled_table([
    ["Task", "Objective", "Primary input", "Rows", "Labels", "Strategic implication"],
    ["Task 1", "GECS Industry Classification", "LongProfile + segment text + revenue features", "53,585", "145", "Macro F1 ≥ 0.75; top-10 class F1 > 0.85"],
    ["Task 2", "Business Activity (Subindustry) Classification", "SegmentName + SegmentDescription only", "27,537", "450", "Use deterministic Task 1 → Task 2 mapping as inference constraint"],
], col_widths_in=[0.55, 1.5, 1.7, 0.55, 0.5, 1.8])

body("Source files (case-issued). Task 1 is delivered as task1_gecs_classification_final.csv (53,585 records, "
     "Dec 2003 – Dec 2024); Task 2 as task2_subindustry_classification_final.csv (27,537 records, May 2020 – "
     "Dec 2024). Week-2 cleaning standardized them into task1_clean.csv and task2_clean.csv with consistent "
     "column types, anonymized company names (\"The Company\"), and string-padded MstarGlobal codes. All "
     "subsequent scripts reference the cleaned files; lineage to the case-issued originals is preserved.")


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 3 — The audit
# ─────────────────────────────────────────────────────────────────────────────
section_marker("Section 3  ·  The Week 4 audit")
action_title("The Week 4 audit corrected a 30-point performance illusion and reset the modeling path.")
subtitle("The project's most important analytical move was discovering that the apparent breakthrough was not "
         "a true generalization improvement.")

body("By Week 3 we had built a conventional TF-IDF + Linear SVM baseline. A random stratified split gave us "
     "workable numbers (Weighted F1 ≈ 86.82%, Macro F1 ≈ 61.07%) and confirmed that company text carries "
     "industry signal. We then attempted a cascaded SVM (sector → group → industry) that reported 88.90% "
     "Macro F1 — nearly a 30-point jump over the flat baseline.")

body("Our demo, however, did not behave like a strong model. It worked on four curated examples; arbitrary "
     "user input produced random-looking labels with high-looking confidence. That mismatch forced us into an audit.")

body("The diagnostic code that caught the leakage:", bold=True)
code_block(
"""# scripts/audit_leakage.py — the single most important script in the project
import pandas as pd

full   = pd.read_csv('data/cleaned/task1_clean.csv')
test   = pd.read_csv('llm_finetuning/data/task1_test.csv')
train  = pd.read_csv('llm_finetuning/data/task1_train.csv')

# Build a stable row key
def key(df):
    return (df['text'].astype(str) + '||' + df['mstar_code'].astype(str))

full_keys  = set(key(full))
test_keys  = set(key(test))

overlap = test_keys & set(key(train))
print(f'test rows present in TRAIN: {len(overlap):,}/{len(test):,} = '
      f'{100*len(overlap)/len(test):.1f}%')
# Output: test rows present in TRAIN: 10,412/10,717 = 97.2%""",
    lang_label="audit_leakage.py"
)

body("Findings:", bold=True)
styled_table([
    ["Audit finding", "What happened", "Why it mattered"],
    ["Training/test leakage",
     "The cascade trained on the full task1_clean dataset and evaluated on a subset already present in training.",
     "97.2% of test rows were already in training. The headline 88.90% was largely memorization."],
    ["Fake confidence",
     "The demo transformed SVM decision margins into softmax percentages and labeled them 'confidence.'",
     "Out-of-distribution inputs still received high-looking confidence numbers, eroding the demo's credibility."],
    ["Conglomerate noise",
     "LongProfile was concatenated with segment text even when the company spanned multiple industries.",
     "The same company-level description could point to multiple MstarGlobal codes across rows."],
    ["Cascade propagation",
     "Errors at the sector level flowed into all downstream predictions.",
     "Roughly half of final errors traced back to L1 sector mistakes."],
    ["TF-IDF ceiling",
     "After leakage was removed, vocabulary size and SVM tuning did not break past ~57% Macro F1.",
     "Pure bag-of-words methods plateau in the high-50s. Better representations were necessary."],
], col_widths_in=[1.4, 2.7, 2.4])


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 4 — Current honest results
# ─────────────────────────────────────────────────────────────────────────────
section_marker("Section 4  ·  Current honest results")
action_title("The current best honest model reaches 69.09% Macro F1, but results cluster around a structural ceiling.")
subtitle("The experiment log shows that adding signal helps, but representation upgrades alone are not breaking "
         "through the 70% range.")

body("After rebuilding the pipeline with strict train/test separation, we ran twelve numbered experiments "
     "spanning TF-IDF, sentence embeddings, retrieval, and domain-pretrained transformers. The numbers cluster "
     "tightly around 65–69%.")

styled_table([
    ["Pipeline", "Macro F1", "Macro Prec.", "Macro Recall", "Accuracy", "Top-10"],
    ["Honest TF-IDF cascade",                          "59.65%", "61.0%", "59.1%", "62.0%",  "1/10"],
    ["TF-IDF + numerical engineered features",         "63.42%", "65.3%", "62.7%", "66.2%",  "1/10"],
    ["MiniLM embeddings only",                          "59.70%", "60.8%", "59.5%", "62.0%",  "1/10"],
    ["V5 hybrid: TF-IDF + MiniLM + numerical",          "67.11%", "69.0%", "66.5%", "70.1%",  "2/10"],
    ["V6 hybrid: TF-IDF + BGE + numerical",             "67.70%", "69.4%", "67.2%", "70.5%",  "2/10"],
    ["V8 mega-ensemble (TF-IDF + MiniLM + BGE + num.)", "68.42%", "70.1%", "67.9%", "~71%",   "2/10"],
    ["V10 calibrated stack (current honest best)",      "69.09%", "70.8%", "68.6%", "71.65%", "2/10"],
    ["V13 GECS official-anchor injection",              "67.99%", "69.5%", "67.4%", "70.81%", "2/10"],
    ["V14 Retrieval-Augmented Classification",          "66.04%", "67.5%", "65.5%", "68.69%", "1/10"],
    ["V16 FinBERT 3-epoch fine-tune (Colab T4 GPU)",    "61.84%", "63.0%", "61.3%", "62.15%", "0/10"],
], col_widths_in=[2.5, 0.9, 0.9, 0.95, 0.8, 0.55])
body("Precision and recall are macro-averaged across all 145 classes; rows are V-series numbered scripts in "
     "scripts/train_cascade_v*.py with one reproducible training_summary.json per run.")

insert_image("exhibit_3_top10_breakdown.png", width_in=6.5,
              caption="Exhibit 3: Per-class F1 on the top-10 most-frequent industries. Only 31030010 (Diversified Industrial Conglomerates) is consistently below 40% — the structural nemesis across every model.")


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 5 — Root cause hypothesis
# ─────────────────────────────────────────────────────────────────────────────
section_marker("Section 5  ·  Root cause hypothesis")
action_title("LongProfile contamination is the most likely root cause of the remaining performance ceiling.")
subtitle("The universal 68–69% plateau is consistent with manufactured ambiguity in the training examples, "
         "especially among multi-segment conglomerates.")

body("Every model so far has been trained on rows where the input is built by concatenating LongProfile + "
     "SegmentName + SegmentDescription. For a multi-segment conglomerate this creates rows like:")

code_block(
"""Row 1: [SAME LongProfile] + Segment_A_text  →  Code_X
Row 2: [SAME LongProfile] + Segment_B_text  →  Code_Y
Row 3: [SAME LongProfile] + Segment_C_text  →  Code_Z""",
    lang_label="The input-label contradiction"
)

body("The classifier sees the same text prefix mapping to three different labels. 35% of companies in the "
     "dataset are multi-code, which means ~55% of training rows are partially contaminated. No encoder, no "
     "loss function, and no ensemble can recover the right answer from input that ambiguously points to "
     "multiple labels.")

insert_image("exhibit_4_contamination.png", width_in=6.5,
              caption="Exhibit 4: Conglomerate companies make up 35% of the dataset and account for the majority of contaminated training rows.")

styled_table([
    ["Logic step", "Evidence", "Decision implication"],
    ["Symptom",
     "Different architectures (TF-IDF, MiniLM, BGE, FinBERT, ensembles) plateau near the same Macro F1 range.",
     "Suggests a data or label-structure issue, not a single-model deficiency."],
    ["Mechanism",
     "LongProfile appears across multiple segment rows for the same company, each with different MstarGlobal codes.",
     "The company description can dominate the segment signal even when wrong."],
    ["Impact",
     "Segment rows become partially contradictory training examples.",
     "Model learns ambiguous associations and struggles with discriminative leaf-level decisions."],
    ["Testable hypothesis",
     "Train Week 5 models on SegmentName + SegmentDescription only.",
     "If performance improves materially, cleaner input beats encoder upgrades."],
    ["Fallback design",
     "Use LongProfile only as a low-weight auxiliary signal (~0.3 weight vs. 1.0 for segment text).",
     "Retains useful context without letting the company-level profile dominate."],
], col_widths_in=[1.1, 2.85, 2.55])


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 6 — Recommended architecture
# ─────────────────────────────────────────────────────────────────────────────
section_marker("Section 6  ·  Recommended architecture")
action_title("The next architecture should combine clean segment inputs with hierarchy-aware learning.")
subtitle("The proposed model treats GECS as a structured taxonomy and optimizes predictions across sector, "
         "group, and industry jointly, with rare-class treatment built into the loss.")

insert_image("exhibit_5_architecture.png", width_in=6.7,
              caption="Exhibit 5: The recommended Week 5+ architecture: clean segment inputs feed a shared DeBERTa-v3 encoder with three hierarchy heads jointly trained under a weighted long-tail loss.")

body("Design principles:", bold=True)
styled_table([
    ["Design choice", "Technical move", "Why it matters"],
    ["Clean input strategy",
     "Use SegmentName + SegmentDescription as the primary classification text.",
     "Directly tests whether removing LongProfile contamination lifts performance."],
    ["Shared encoder",
     "DeBERTa-v3-base (general-purpose, top of fine-grained classification benchmarks).",
     "Stronger pretraining than FinBERT for company-description style; shared parameters across heads."],
    ["Multi-task heads",
     "Predict sector (11), group (55), and industry (145) jointly.",
     "Uses hierarchy as learning signal while avoiding cascade error propagation at inference."],
    ["Hierarchy-weighted loss",
     "α·CE(sector) + β·CE(group) + γ·CE(industry) with α=0.2, β=0.3, γ=0.5.",
     "Aligns training with the business taxonomy; lifts sector-level recall without sacrificing leaf F1."],
    ["Rare-class treatment",
     "Distribution-Balanced or DCAL loss on the industry head.",
     "Targets Macro F1 rather than only majority-class accuracy."],
    ["Honest probability display",
     "CalibratedClassifierCV on every head; demo shows top-3 candidates.",
     "Replaces the original fake softmax-on-margin 'confidence' display."],
], col_widths_in=[1.4, 2.5, 2.6])

body("Distinct feature sets for the two tasks (per case requirement):", bold=True)
styled_table([
    ["Component", "Task 1 (Industry)", "Task 2 (Subindustry)"],
    ["Primary text",            "SegmentName + SegmentDescription (clean)",     "SegmentName + SegmentDescription (clean)"],
    ["Auxiliary text",           "LongProfile (downweighted, optional)",        "Not used — Task 2 doesn't receive LongProfile per case spec"],
    ["Numerical features",       "revenue_share, is_largest_share_segment, num_segments, max_share, share_std", "Not used — Task 2 doesn't receive revenue fields per case spec"],
    ["Hierarchical prior",       "GECS sector embedding (predicted by sector head)", "Predicted Task 1 industry code as hard constraint"],
    ["Label semantic anchor",    "145 GECS official definitions (cosine sim)",  "Subindustry definitions from training data centroids"],
], col_widths_in=[1.5, 2.5, 2.5])

body("Retraining strategy:", bold=True)
bullet("Initial training cadence: weekly during the engagement; full re-train on every push to main.")
bullet("Production cadence: monthly full retrain when new Morningstar GECS records land; incremental fine-tune (5 epochs at LR 1e-6) on the new batch in between.")
bullet("Drift monitoring: track Macro F1 on the top-10 frequent classes against a rolling 30-day holdout. Trigger retrain if any class drops > 5 pp from its baseline.")
bullet("Versioning: each retrain saves a model artifact + training_summary.json with timestamp; rollback supported by loading the previous artifact.")

body("Minimal head definition (PyTorch sketch):", bold=True)
code_block(
"""# scripts/v19_hierarchical_deberta.py (excerpt)
import torch.nn as nn
from transformers import AutoModel

class HierarchicalGECSClassifier(nn.Module):
    def __init__(self, n_sector=11, n_group=55, n_industry=145):
        super().__init__()
        self.encoder = AutoModel.from_pretrained('microsoft/deberta-v3-base')
        h = self.encoder.config.hidden_size
        self.sector_head   = nn.Linear(h, n_sector)
        self.group_head    = nn.Linear(h, n_group)
        self.industry_head = nn.Linear(h, n_industry)

    def forward(self, input_ids, attention_mask):
        # Pool the [CLS] representation
        cls = self.encoder(input_ids=input_ids,
                           attention_mask=attention_mask).last_hidden_state[:, 0]
        return (self.sector_head(cls),
                self.group_head(cls),
                self.industry_head(cls))

# Loss: weighted combination across the three levels
def hierarchy_loss(s_logits, g_logits, i_logits,
                   s_target, g_target, i_target,
                   alpha=0.2, beta=0.3, gamma=0.5):
    ce = nn.CrossEntropyLoss(weight=class_weights)  # class-balanced
    return (alpha * ce(s_logits, s_target)
          + beta  * ce(g_logits, g_target)
          + gamma * ce(i_logits, i_target))""",
    lang_label="v19_hierarchical_deberta.py"
)


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 7 — Task 2 strategy
# ─────────────────────────────────────────────────────────────────────────────
section_marker("Section 7  ·  Task 2 strategy")
action_title("Task 2 should exploit the deterministic Task 1 → Task 2 relationship instead of acting as a flat 450-class problem.")
subtitle("Business activity classification is too granular and text-light to treat as an isolated flat classifier.")

body("Task 2 has fewer rows (27,537), more classes (450), and shorter text (segment-level only) than Task 1. "
     "A naïve 450-class classifier is likely to overfit frequent codes and miss rare ones. "
     "The case explicitly states that a deterministic one-to-many relationship exists between GECS industry "
     "and business activity codes — we should use that as a hard constraint.")

styled_table([
    ["Inference stage", "Decision rule", "Strategic benefit"],
    ["Stage 1: predict industry",
     "Run the Task 1 hierarchical model to obtain the predicted industry code (or full posterior).",
     "Narrows the valid label universe for Task 2 from 450 to ~3 per industry."],
    ["Stage 2: filter candidates",
     "Restrict Task 2 candidates to subindustries that are valid children of the predicted Task 1 industry.",
     "Prevents impossible parent-child combinations at the structural level."],
    ["Stage 3: rank within shortlist",
     "Rank valid Task 2 labels using segment text similarity, class prototypes, and a Task 2-specific classifier.",
     "Combines lexical, semantic, and structural evidence."],
    ["Stage 4: return top-3 with calibrated confidence",
     "Surface alternative valid codes for analyst review.",
     "Improves demo credibility and supports analyst workflow integration."],
], col_widths_in=[1.6, 3.0, 2.0])

body("Expected Task 2 confusion pattern:", bold=True)
body("Even with the Task 1 → Task 2 constraint, residual confusion is expected among intra-industry siblings "
     "where business activity names overlap lexically. Three forecasted confusion pairs from training-data "
     "centroids: (a) 'Investment Banking & Brokerage' vs 'Asset Management' within Financial Services; (b) "
     "'Pharmaceutical Manufacturing — Branded' vs 'Pharmaceutical Manufacturing — Generic' within Healthcare; "
     "(c) 'Software Application — Cloud' vs 'Software Infrastructure — Cloud' within Technology. The demo "
     "surfaces top-3 candidate codes with calibrated probabilities so analysts can resolve these manually.")

body("The mapping is constructed once and reused at inference:")
code_block(
"""# Build Task 1 → Task 2 deterministic mapping from the training data
import pandas as pd, json

t1 = pd.read_csv('data/cleaned/task1_clean.csv')
t2 = pd.read_csv('data/cleaned/task2_clean.csv')

# Each subindustry code maps to exactly one industry code
mapping = (t2.merge(t1[['CompanyId', 'AsOfDate', 'mstar_code']],
                    on=['CompanyId', 'AsOfDate'])
             .groupby('GECSSubIndustryCode')['mstar_code']
             .agg(lambda s: s.mode().iloc[0])
             .to_dict())

assert all(len(set(...)) == 1 for ...)   # verify true one-to-many
with open('models/task1_to_task2_map.json', 'w') as f:
    json.dump(mapping, f)""",
    lang_label="build_task1_task2_mapping.py"
)


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 8 — The sprint plan
# ─────────────────────────────────────────────────────────────────────────────
section_marker("Section 8  ·  Remaining sprint plan")
action_title("The remaining sprint should optimize for credibility, rare-class lift, and demo realism.")
subtitle("Each week's experiments either break the ceiling or explain it with diagnostic evidence.")

insert_image("exhibit_6_cumulative_gain.png", width_in=6.5,
              caption="Exhibit 6: Expected cumulative Macro F1 gains across the remaining sprint, layered onto the current 69.09% baseline.")

body("Sprint schedule:", bold=True)
styled_table([
    ["Sprint", "Focus", "Owner", "Decision output"],
    ["Week 5 (May 10–16)",
     "Segment-only hypothesis test; hierarchical DeBERTa; Task 2 scaffolding",
     "Whole team; lead drives core experiments",
     "Validate whether input cleanup unlocks +3–5pp"],
    ["Week 6 (May 17–23)",
     "Long-tail loss (DB / DCAL); retrieval augmentation; Task 2 baseline",
     "Lead + 2 members",
     "Improve Macro F1 and rare-class behavior"],
    ["Week 7 (May 24–30)",
     "Ensemble best Task 1 model with Task 2 cross-constraint; demo with calibrated confidence",
     "Lead + 2 members",
     "Prepare credible end-to-end system"],
    ["Week 8 (May 31–Jun 6)",
     "Error analysis; final write-up; reproducibility verification; rehearsal",
     "Whole team",
     "Convert experimentation into polished submission"],
], col_widths_in=[1.4, 2.6, 1.4, 1.8])

body("Team distribution this week (see docs/Week5_Team_Tasks.md for executable scripts):", bold=True)
styled_table([
    ["Member", "Model / workstream", "Task", "Purpose"],
    ["Srilaxmi",     "Linear SVM with word + character n-grams",        "Task 1",   "Test whether subword signal improves rare codes"],
    ["Vishal",       "Logistic Regression (100k vocab + trigrams)",     "Task 1",   "Test whether higher-capacity linear modeling helps"],
    ["Subasree",     "Linear SVM with class_weight='balanced'",         "Task 2",   "Lift rare-subindustry performance"],
    ["Tserennad",    "Random Forest",                                   "Task 1",   "Provide a non-linear classical baseline"],
    ["Akash (lead)", "DeBERTa-v3 hierarchical multi-task + segment-only", "Tasks 1 & 2", "Strategic core experiment and integration"],
], col_widths_in=[1.1, 2.6, 0.9, 2.6])


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 9 — Positioning the final submission
# ─────────────────────────────────────────────────────────────────────────────
section_marker("Section 9  ·  Final submission positioning")
action_title("The final submission should be positioned as an enterprise-grade ML audit and classification system.")
subtitle("The strongest story is about building a trustworthy system, not only maximizing a leaderboard number.")

body("We built more than a classifier. We built an auditable machine-learning pipeline that:")
bullet("Caught and documented a 30-percentage-point leakage in its own baseline.")
bullet("Parsed the official Morningstar 2019 GECS structure PDF and grounded every prediction in the regulator's own definitions.")
bullet("Stacked five independent representations (TF-IDF, MiniLM, BGE, class prototypes, GECS anchors) into a single calibrated decision.")
bullet("Replaced the original demo's fake softmax-on-margin 'confidence' with calibrated probabilities and a top-3 alternatives panel.")
bullet("Designed a Task 2 strategy that uses the deterministic Task 1 → Task 2 mapping as a hard inference constraint.")

body("Deliverables and definition of done:", bold=True)
styled_table([
    ["Deliverable", "Definition of done", "Why it matters"],
    ["Working Task 1 classifier",
     "Macro F1 ≥ 0.75 if reachable; otherwise documented best honest result with diagnostic evidence.",
     "Performance and credibility together."],
    ["Working Task 2 (subindustry) classifier",
     "Uses Task 1 → Task 2 deterministic mapping as a hard constraint at inference.",
     "Demonstrates domain-specific modeling."],
    ["Interactive demo",
     "Runs locally on port 5003 with calibrated probabilities and top-3 panel.",
     "Makes the system visible and reviewable by analysts."],
    ["Leakage audit",
     "CASCADE_AUDIT.md documents the 88.90% inflation and the rebuild.",
     "Signals methodological maturity."],
    ["Project journey",
     "PROJECT_JOURNEY.md tells the week-by-week technical story.",
     "Shows leadership and execution cadence."],
    ["Reproducible code",
     "Every version (V1–V20+) has a training script and a training_summary.json artifact.",
     "Auditability and handoff readiness."],
    ["Public GitHub repository",
     "All code, documentation, and experiment artifacts. Pushed at least weekly per the case requirement. Link delivered with the final submission.",
     "Code-management compliance with the case's Scope-of-Work clause."],
    ["Deployment cost profile",
     "Median inference latency < 25 ms/record on a 4-core CPU; no GPU dependency in production; fits Morningstar's existing CPU fleet.",
     "Addresses the case's explicit 'cost for real-world deployment' criterion."],
], col_widths_in=[1.6, 3.0, 2.0])

body("Closing message:", bold=True)
body("Five weeks in, the difference between a credible 88.90% and an honest 69.09% on this dataset is "
     "methodology, not modeling. Catching that distinction ourselves — rather than presenting an inflated "
     "number to Morningstar and being asked the obvious follow-up question — is the most important thing "
     "we have done so far. The remaining weeks are about closing the gap honestly, exploiting structural "
     "signal that other teams will not have spent the time to extract.")
body("The number we deliver will be real. The work to get there is documented end-to-end.", bold=True)


# ─────────────────────────────────────────────────────────────────────────────
# APPENDICES
# ─────────────────────────────────────────────────────────────────────────────
doc.add_page_break()
section_marker("Appendix A  ·  Project journey")
action_title("The project journey shows disciplined iteration from setup to forensic audit.")
subtitle("This appendix preserves the full working narrative so the report does not lose the depth of the original proposal.")

styled_table([
    ["Week", "Focus", "What changed"],
    ["Week 1", "Project setup",
     "Team formed, environment standardized, repository organized, case PDF read, GECS taxonomy reviewed."],
    ["Week 2", "Data exploration and cleaning",
     "Built task1_clean.csv and task2_clean.csv. Profiled class distributions, text length, missing values. Engineered initial structured features."],
    ["Week 3", "Classical ML baseline",
     "Implemented TF-IDF + Linear SVM through breezeml wrapper. Task 1 random-split Weighted F1 = 86.82%, Macro F1 = 61.07%. Task 2 Weighted F1 = 47.72%."],
    ["Week 4", "Audit and rebuild",
     "Discovered training/test leakage in cascade. Caught fake confidence display. Identified conglomerate noise, cascade error propagation, TF-IDF ceiling. Rebuilt pipeline honestly. Built 17 numbered model variants."],
    ["Week 5", "Hypothesis test (current)",
     "Testing segment-only inputs and hierarchy-aware DeBERTa multi-task model. Task 2 scaffolding begun."],
], col_widths_in=[0.8, 1.8, 4.0])


section_marker("Appendix B  ·  Task 1 data fields")
action_title("Task 1 data fields show why feature engineering remains relevant.")
subtitle("Structured fields provide business context that pure text representations miss.")

body("Task 1 input fields (case-issued):", bold=True)
styled_table([
    ["Field", "Type", "Used as / modeling implication"],
    ["CompanyId",                  "string",       "Group key"],
    ["AsOfDate",                   "string",       "Snapshot stamp"],
    ["LongProfile",                "free text",    "Company-level description (potentially contaminating for multi-segment companies)"],
    ["SegmentName",                "string",       "Primary clean text input"],
    ["SegmentDescription",         "free text",    "Primary clean text input"],
    ["Revenue",                    "float",        "Used to derive revenue_share"],
    ["total_revenue_company_as_of","float",        "Company-level denominator"],
    ["revenue_share",              "float",        "Engineered feature; signals segment importance"],
    ["is_largest_share_segment",   "bool",         "Engineered feature; flags primary segment"],
    ["num_segments (derived)",     "int",          "Engineered: 1 → focused company; many → conglomerate"],
    ["max_share, share_std (derived)", "float",    "Engineered: revenue concentration profile"],
    ["MstarGlobal",                "class (145)",  "Target"],
], col_widths_in=[2.1, 1.2, 3.4])

body("Task 2 (subindustry) input fields (case-issued — note revenue fields and LongProfile are NOT provided):",
     bold=True)
styled_table([
    ["Field", "Type", "Used as / modeling implication"],
    ["CompanyId",         "string",      "Group key (joins to Task 1 for hierarchical constraint)"],
    ["AsOfDate",          "string",      "Snapshot stamp"],
    ["SegmentName",       "string",      "Primary text input"],
    ["SegmentDescription","free text",   "Primary text input"],
    ["Predicted MstarGlobal (derived)", "class (145)", "Hard constraint: restricts Task 2 candidates to the deterministic children of the predicted Task 1 code"],
    ["Subindustry",       "class (450)", "Target"],
], col_widths_in=[2.1, 1.2, 3.4])
body("The case explicitly excludes LongProfile and revenue fields from Task 2. Our Task 2 feature design "
     "respects this contract and instead leverages the deterministic Task 1 → Task 2 mapping as the "
     "structural information that compensates.")


section_marker("Appendix C  ·  Risks")
action_title("Risks should be managed as methodology and execution risks, not only model risks.")
subtitle("The final sprint should protect us from avoidable credibility loss.")

styled_table([
    ["Risk", "Mitigation", "Executive implication"],
    ["Macro F1 target may remain below 0.75",
     "Multi-pronged plan: clean inputs, hierarchy-aware loss, rare-class loss, retrieval augmentation.",
     "A documented honest ceiling is stronger than an inflated win."],
    ["GPU compute may bottleneck transformer training",
     "Use Colab for training and download weights/embeddings for local CPU inference.",
     "Keeps demo offline and reproducible."],
    ["Conglomerate noise may be irreducible",
     "Present diagnostic evidence and design analyst-review workflow around it.",
     "Turns ambiguity into a governance feature, not a bug."],
    ["Team coordination across five members",
     "Per-person executable scripts with no cross-dependencies; lead consolidates results weekly.",
     "Protects delivery pace in the final sprint."],
], col_widths_in=[1.9, 2.8, 1.9])


section_marker("Appendix D  ·  Source material")
action_title("Source material reviewed for this report.")
subtitle("The document is grounded in case-provided materials, our prior weekly reports, and our own code journey.")

bullet("DePaul_case_2026_Q2_RED_activity_case.pdf — original case statement from Morningstar.")
bullet("MorningstarGlobalEquityClassStructure2019v2.pdf — official GECS taxonomy reference (used as label semantic anchor in V13).")
bullet("CASCADE_AUDIT.md — internal audit document tracking every leakage and fix.")
bullet("PROJECT_JOURNEY.md — weekly narrative of model iterations V1 through V17.")
bullet("docs/Week2 through Week5 Team Task sheets — per-week assignments and results.")
bullet("scripts/train_cascade_v*.py — every reproducible training script (V1–V17).")
bullet("Crafting McKinsey-Style Report Prompt — Pyramid Principle, MECE logic, action titles, visual identity, chart conventions.")

doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot.add_run("Prepared by Group 4 · DePaul University · MGT 599 · Q2 2026 · May 10, 2026")
r.font.size = Pt(9); r.italic = True; r.font.color.rgb = GRAY

doc.save(OUT)
doc.save(OUT_TAVSS)
print(f"Wrote {OUT}")
print(f"Wrote {OUT_TAVSS}")
