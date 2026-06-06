"""Convert Week5_Team_Tasks.md to a clean Word doc."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
OUT  = ROOT / "docs" / "Week5_Team_Tasks.docx"

doc = Document()
# Default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)


def para(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def code_block(code: str):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    # Light gray background via paragraph shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p


def hrule():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Title block ──
title = doc.add_heading("Week 5 — Team Task Sheet", level=0)
para("MGT 599 Capstone · Group 4 · DePaul University Chicago", italic=True)

para("Each person runs one independent model this week. No coordination needed — just run your script, "
     "record your numbers, and share results with Akash.")
para("Do NOT touch server_legendary.py, train_cascade.py, or anything in legendary/ — that is Akash's work.")
hrule()

# ── Setup ──
heading("Setup (everyone does this first)", level=2)
code_block(
'cd "C:\\Users\\akash\\Desktop\\capstone MGT 599"\n'
"pip install breezeml==0.2.5 lightgbm scikit-learn\n"
'python -c "import breezeml; print(breezeml.__version__)"  # should print 0.2.5'
)
hrule()

# ── Srilaxmi ──
heading("Srilaxmi — Linear SVM with character n-grams on Task 1", level=2)
code_block(
"""# save as: scripts/srilaxmi_week5.py
# run:     python scripts/srilaxmi_week5.py

import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.model_selection import train_test_split
from scipy.sparse import hstack
from breezeml import classifiers

df = pd.read_csv("data/cleaned/task1_clean.csv")
df = df.dropna(subset=["text", "mstar_code"])

# Word n-grams
vec_word = TfidfVectorizer(max_features=50_000, ngram_range=(1, 2), sublinear_tf=True)
X_word   = vec_word.fit_transform(df["text"].tolist())

# Character n-grams (new this week)
vec_char = TfidfVectorizer(max_features=30_000, ngram_range=(3, 5),
                            analyzer="char_wb", sublinear_tf=True)
X_char   = vec_char.fit_transform(df["text"].tolist())

X = hstack([X_word, X_char]).tocsr()
y = df["mstar_code"].astype(str).tolist()

X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.2, stratify=y, random_state=42
)

_, report = classifiers.linear_svm(X=X_train, y=y_train, X_test=X_test, y_test=y_test)
print(report)"""
)
para("Record your score:", bold=True)
para("• Macro F1: ______")
para("• Weighted F1: ______")
para("• Last week's number to beat: Macro F1 = (your Week 4 number)")
hrule()

# ── Vishal ──
heading("Vishal — Logistic Regression with larger vocabulary on Task 1", level=2)
code_block(
"""# save as: scripts/vishal_week5.py
# run:     python scripts/vishal_week5.py

import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.model_selection import train_test_split
from breezeml import classifiers

df = pd.read_csv("data/cleaned/task1_clean.csv")
df = df.dropna(subset=["text", "mstar_code"])

# Bigger vocabulary + trigrams this week
vec = TfidfVectorizer(max_features=100_000, ngram_range=(1, 3),
                       sublinear_tf=True, min_df=2)
X   = vec.fit_transform(df["text"].tolist())
y   = df["mstar_code"].astype(str).tolist()

X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.2, stratify=y, random_state=42
)

_, report = classifiers.logistic_regression(
    X=X_train, y=y_train, X_test=X_test, y_test=y_test
)
print(report)"""
)
para("Record your score:", bold=True)
para("• Macro F1: ______")
para("• Weighted F1: ______")
para("• Last week's number to beat: Macro F1 = (your Week 4 number)")
hrule()

# ── Subasree ──
heading("Subasree — Linear SVM with balanced class weights on Task 2", level=2)
code_block(
"""# save as: scripts/subasree_week5.py
# run:     python scripts/subasree_week5.py

import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.model_selection import train_test_split
from breezeml import classifiers

df = pd.read_csv("data/cleaned/task2_clean.csv")
df = df.dropna(subset=["text"])

label_col = "GECSSubIndustryCode"   # update if your column name is different
y = df[label_col].astype(str).tolist()

# Larger vocab this week
vec = TfidfVectorizer(max_features=30_000, ngram_range=(1, 2),
                       sublinear_tf=True, min_df=2)
X   = vec.fit_transform(df["text"].tolist())

X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.2, stratify=y, random_state=42
)

_, report = classifiers.linear_svm(
    X=X_train, y=y_train, X_test=X_test, y_test=y_test,
    class_weight="balanced"   # rebalance for rare codes
)
print(report)"""
)
para("Record your score:", bold=True)
para("• Macro F1: ______")
para("• Weighted F1: ______")
para("• Last week's number to beat: Weighted F1 = (your Week 4 number)")
hrule()

# ── Tserennad ──
heading("Tserennad — Random Forest on Task 1", level=2)
code_block(
"""# save as: scripts/tserennad_week5.py
# run:     python scripts/tserennad_week5.py

import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.model_selection import train_test_split
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import classification_report

df = pd.read_csv("data/cleaned/task1_clean.csv")
df = df.dropna(subset=["text", "mstar_code"])

vec = TfidfVectorizer(max_features=20_000, ngram_range=(1, 2), sublinear_tf=True)
X   = vec.fit_transform(df["text"].tolist())
y   = df["mstar_code"].astype(str).tolist()

X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.2, stratify=y, random_state=42
)

clf = RandomForestClassifier(
    n_estimators=300,
    max_depth=40,
    min_samples_leaf=2,
    class_weight="balanced",
    n_jobs=-1,
    random_state=42,
)
clf.fit(X_train, y_train)
preds = clf.predict(X_test)
print(classification_report(y_test, preds, zero_division=0, digits=4))"""
)
para("Record your score:", bold=True)
para("• Macro F1: ______")
para("• Weighted F1: ______")
para("• Last week's number to beat: Macro F1 = (your Week 4 number)")
hrule()

# ── Results table ──
heading("Results Table — fill in and send to Akash", level=2)

table = doc.add_table(rows=6, cols=5)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Person"
hdr[1].text = "Model"
hdr[2].text = "Task"
hdr[3].text = "Macro F1"
hdr[4].text = "Weighted F1"

rows = [
    ("Srilaxmi",  "Linear SVM (word + char n-grams)",      "Task 1", "", ""),
    ("Vishal",    "Logistic Regression (100k, trigrams)",  "Task 1", "", ""),
    ("Subasree",  "Linear SVM (balanced weights)",         "Task 2", "", ""),
    ("Tserennad", "Random Forest",                         "Task 1", "", ""),
    ("Best from Week 4", "(whatever you scored)",          "",       "", ""),
]
for i, row in enumerate(rows, start=1):
    cells = table.rows[i].cells
    for j, val in enumerate(row):
        cells[j].text = val

# Bold header row
for cell in table.rows[0].cells:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

doc.add_paragraph()
para("MGT 599 Capstone · Group 4 · DePaul University Chicago · Week 5 · May 2026",
     italic=True, size=9)

doc.save(OUT)
print(f"Wrote {OUT}")
