"""Render Week5_Report.md as a clean Word doc matching the Week 3/4 report style."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
OUT  = ROOT / "docs" / "Week5_Report.docx"

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
TEAL = RGBColor(0x0E, 0x6B, 0x6E)
GRAY = RGBColor(0x3D, 0x49, 0x51)

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.2); s.bottom_margin = Cm(2.2)
    s.left_margin = Cm(2.0); s.right_margin = Cm(2.0)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def title(text, size=20, color=NAVY, bold=True):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.size = Pt(size); r.font.color.rgb = color; r.bold = bold


def heading(text, level=2):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = NAVY
        r.font.size = Pt(14 if level == 2 else 12)


def para(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    r.font.color.rgb = GRAY


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text); r.font.size = Pt(11); r.font.color.rgb = GRAY


def shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def styled_table(rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]; run = p.add_run(str(val))
            run.font.size = Pt(10); run.font.color.rgb = GRAY
            if r_idx == 0:
                run.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                shade_cell(cell, "1F3A5F")
            elif r_idx % 2 == 1:
                shade_cell(cell, "F7F7F7")
    doc.add_paragraph()


def code_block(code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(8); p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "F4F4F0")
    pPr.append(shd)
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left"); left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24"); left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "0E6B6E"); pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(code); r.font.name = "Consolas"
    r.font.size = Pt(9); r.font.color.rgb = GRAY


def hrule():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom"); bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6"); bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto"); pBdr.append(bottom)
    pPr.append(pBdr)


# ── Title ──
title("MGT 599 Capstone — Week 5 Report", size=22)
para("Group 4 | Akash Anipakalu Giridhar", bold=True)
para("Date: May 10, 2026", italic=True)
hrule()


# ── 1. Description of Work ──
heading("1. Description of Work")

heading("What We Were Trying to Achieve", level=3)
para("The goal for Week 5 was to break past the ~68–69% Macro F1 ceiling that every model — TF-IDF, MiniLM, "
     "BGE, FinBERT, multi-encoder ensembles — converged to in Week 4. The diagnostic hypothesis from the Week 4 "
     "audit was clear:")
para("The universal plateau is not a representation problem. It is an input-contamination problem. "
     "LongProfile concatenated with segment text creates many-to-many label mappings for the 35% of companies "
     "that are diversified conglomerates.", italic=True)
para("Three concrete questions guided the week:")
bullet("Does removing LongProfile from the input unlock meaningful F1 gains? (Lane A: segment-only baseline.)")
bullet("Does a hierarchy-aware transformer trained jointly on sector → group → industry beat a flat 145-class head? (Strategic lane: DeBERTa-v3 multi-task on Colab GPU.)")
bullet("Does retrieval-augmented classification using the official GECS taxonomy as anchors add usable signal? (V13 anchor injection completed.)")

heading("Approach", level=3)
para("The week split into two parallel tracks. The classical-ML team ran four model variants on the existing "
     "concatenated input to establish whether smaller architectural changes (subword features, larger vocab, "
     "class rebalancing, non-linear classifiers) could close the gap. The strategic track moved to Google Colab "
     "GPU for the first transformer fine-tune the project has attempted at scale.")

para("Team distribution (executable scripts in docs/Week5_Team_Tasks.md):", bold=True)
styled_table([
    ["Member", "Model", "Task", "Purpose"],
    ["Srilaxmi",     "Linear SVM with word + character n-grams",   "Task 1", "Test whether subword signal helps rare codes"],
    ["Vishal",       "Logistic Regression, 100k vocab + trigrams", "Task 1", "Test whether higher-capacity linear modeling helps"],
    ["Subasree",     "Linear SVM with class_weight='balanced'",    "Task 2", "Lift rare-subindustry F1"],
    ["Tserennad",    "Random Forest, 300 trees, max_depth 40",     "Task 1", "Non-linear classical baseline"],
    ["Akash (lead)", "DeBERTa-v3-base, hierarchical multi-task",   "Task 1 & 2", "Strategic transformer experiment on Colab T4 GPU"],
])

para("Hierarchical multi-task architecture (lead workstream):", bold=True)
para("A single DeBERTa-v3-base encoder feeds three classification heads — sector (11), industry group (55), "
     "industry (145) — trained with a weighted multi-task loss. The hierarchy is treated as a learning signal "
     "rather than a top-down cascade.")
code_block("total_loss = 0.15 · CE(sector) + 0.15 · CE(group) + 0.70 · CE(industry)")

para("Key training configuration:", bold=True)
bullet("Encoder: microsoft/deberta-v3-base (184M parameters)")
bullet("Sequence length: 512 tokens")
bullet("Effective batch size: 64 (per-device 4 with gradient accumulation 16)")
bullet("Optimizer: AdamW, weight decay 0.01")
bullet("Schedule: cosine with linear warmup (5% of total steps)")
bullet("Mixed precision: FP16 autocast with GradScaler")
bullet("Class weighting: sqrt-balanced on the leaf head only")

para("V13 follow-up — GECS official-taxonomy anchoring:", bold=True)
para("Completed the implementation of label semantic anchoring. All 145 official Morningstar 2019 GECS industry "
     "definitions were parsed from the case-issued PDF, encoded with MiniLM and BGE, and added as 580 "
     "cosine-similarity features per sample. Stacked on top of the V8 mega-ensemble, this produced a "
     "123,469-dimensional input for a LinearSVC.")
hrule()


# ── 2. Summary of Findings ──
heading("2. Summary of Findings")

heading("Task 1 — Industry Classification", level=3)
styled_table([
    ["Pipeline", "Macro F1", "Accuracy", "Top-10 pass"],
    ["V10 calibrated stack (Week 4 best)",                "69.09%", "71.65%", "2/10"],
    ["V13 with GECS PDF anchors (Week 5)",                "67.99%", "70.81%", "2/10"],
    ["V14 Retrieval-Augmented Classification",             "66.04%", "68.69%", "1/10"],
    ["DeBERTa-v3 hierarchical (Colab, 2 epochs interim)",  "44.49%", "—",      "—"],
    ["DeBERTa-v3 hierarchical (Colab, after recovery attempts)", "unstable", "—", "—"],
])

para("Key insight 1 — anchors didn't break through.", bold=True)
para("The GECS taxonomy similarity features added measurable signal but were drowned out by the ~122k TF-IDF "
     "features that dominate the input matrix. The official definitions sit at semantic distances from real "
     "company descriptions that differ from the empirical class prototypes we already had. The novel methodology "
     "survives as a documented contribution; the F1 effect was neutral.")

para("Key insight 2 — the DeBERTa fine-tune is mid-flight and unstable.", bold=True)
para("The hierarchical multi-task model achieved 44.49% Macro F1 after two epochs on the original training "
     "settings (focal loss with class weights + weighted sampler + LR 1e-5 + gradient checkpointing). Three "
     "recovery attempts uncovered three real engineering problems with mixed-precision training in Colab's "
     "pinned environment:")
bullet("FP16 gradient unscaling error — GradScaler.unscale_ raised ValueError after the first accumulation step. Root cause: gradient checkpointing with use_reentrant=True caused recomputed backward passes to produce FP16 grads outside the autocast scope. Workable fix: disable gradient checkpointing entirely and halve per-step batch.")
bullet("Double-rebalancing collapse — the first successful pass with WeightedRandomSampler + class_weight='balanced' + focal gamma=2 collapsed the resumed model from 44.49% to 5.45% in one epoch. The compounded rare-class gradients destroyed the representations.")
bullet("Reset to vanilla training — plain shuffle, vanilla CrossEntropyLoss, sqrt-balanced leaf weights, LR 2e-6. Currently re-running with a pre-training sanity-check eval to verify the checkpoint is intact.")

para("Key insight 3 — the V10 calibrated stack remains the project's honest best.", bold=True)
para("Until DeBERTa stabilizes, the production candidate for Task 1 is V10 at 69.09% Macro F1 / 71.65% accuracy "
     "on the case-standard row-level split.")

heading("Task 2 — Subindustry Classification", level=3)
styled_table([
    ["Pipeline", "Macro F1", "Weighted F1"],
    ["Linear SVM with balanced class weights (Subasree, Week 5)", "(pending)", "(pending)"],
    ["Hierarchical roll-up plan from Task 1",                     "scaffolding ready", "scaffolding ready"],
])
para("The Task 2 architecture for the final delivery uses the deterministic Task 1 → Task 2 mapping as a hard "
     "inference constraint: predict the Task 1 industry code, then restrict Task 2 predictions to the ~3 "
     "business-activity codes that roll up into that industry. This avoids treating Task 2 as an isolated "
     "450-class flat problem.")

heading("Broader Insight", level=3)
para("Week 5 confirmed three things and changed our priors on a fourth:")
bullet("Confirmed: the 68–69% ceiling is real for the classical + sentence-embedding family on the contaminated input.")
bullet("Confirmed: hierarchy-aware modeling is the right architectural direction. The DeBERTa multi-task structure is sound; the instability is in training dynamics, not in the design.")
bullet("Confirmed: method documentation matters — every collapse this week was caught and explained.")
bullet("Changed: we expected segment-only inputs to be the dominant lever. The DeBERTa instability prevented a clean test. The hypothesis test moves into early Week 6.")
hrule()


# ── 3. Supporting Outputs ──
heading("3. Supporting Outputs")
styled_table([
    ["File", "Contents"],
    ["colab/finbert_finetune.ipynb",                "Colab notebook for the Week 4 FinBERT baseline (61.84%)"],
    ["scripts/train_cascade_v13_gecs_anchors.py",    "V13 — GECS taxonomy anchor injection + class prototypes"],
    ["scripts/train_cascade_v14_rac.py",             "V14 — Retrieval-Augmented Classification"],
    ["scripts/parse_gecs_taxonomy.py + fill_missing_gecs.py", "PDF parser for all 145 GECS industry definitions"],
    ["gecs_taxonomy.json",                            "Structured taxonomy used by V13"],
    ["models_v13/training_summary.json + models_v14/training_summary.json", "Reproducible results"],
    ["docs/Week5_Team_Tasks.md + .docx",              "Per-member executable scripts and results table"],
    ["docs/Initial_Proposal.md + .docx + _McKinsey.docx", "First formal proposal with Week 1–5 code journey"],
    ["docs/proposal_exhibits/*.png",                  "Six chart exhibits used in the proposal"],
    ["PROJECT_JOURNEY.md, CASCADE_AUDIT.md",          "Running narrative and the leakage audit document"],
])
hrule()


# ── 4. Reflection ──
heading("4. Reflection")

heading("Challenges Encountered", level=3)
para("Challenge 1: Mixed-precision training instability on Colab.", bold=True)
para("Three separate FP16-grad errors surfaced over the recovery attempts (use_reentrant, sampler + weights "
     "interaction, scaler state mismatch on resume). The lesson is that resuming a fine-tuned transformer is "
     "significantly more fragile than initial training — optimizer state, scaler state, and the loss "
     "configuration all have to remain consistent between runs. Future Colab runs will save the full training "
     "state (model + optimizer + scaler + scheduler + best-F1), not weights only.")

para("Challenge 2: GECS anchors got diluted in the feature stack.", bold=True)
para("The 580 anchor features were less than 0.5% of the 123,469-dimensional input. LinearSVC weights them by "
     "gradient contribution, which gave the TF-IDF features structural dominance. The right way to use the "
     "anchor signal is either as a separate scoring branch or by training with a much smaller TF-IDF block and "
     "letting embeddings + anchors dominate. Queued for Week 6.")

para("Challenge 3: The 88.90% leakage diagnosis is right but not yet paid back.", bold=True)
para("We have an honest 69%. Other teams may be presenting higher numbers — almost certainly with the same kind "
     "of contamination we caught in ourselves. There is a real concern that our honesty looks like "
     "underperformance until the methodology audit is read alongside the F1 number. The Week 6–8 plan responds "
     "by making the audit a foreground deliverable (a one-page summary slide).")

para("Challenge 4: My fault, documented.", bold=True)
para("Two of the model collapses this week resulted from changing the wrong hyperparameter at the right time. "
     "When the FP16 error surfaced, the focus shifted to gradient flow plumbing and the loss configuration was "
     "not re-validated for its interaction with the sampler. The post-mortem fix — disabling either the sampler "
     "or class weights but not both — is now codified as a project invariant.")

heading("Next Steps (Week 6)", level=3)
bullet("Finish the DeBERTa recovery run with vanilla settings (plain shuffle, vanilla CE, LR 2e-6, sqrt-balanced leaf weights). Capture the pre-training sanity F1 to verify the loaded checkpoint is intact.")
bullet("Run the segment-only experiment cleanly. If segment-only with the V8 hybrid stack hits 73%+, the contamination hypothesis is confirmed and we re-launch DeBERTa on segment-only inputs.")
bullet("Add Distribution-Balanced or DCAL loss to the leaf head on the next DeBERTa run for explicit rare-class macro F1 lifting.")
bullet("Build Task 2 baseline with the hierarchical roll-up constraint. Target: 50%+ Macro F1.")
bullet("Demo cleanup. Replace server_legendary.py's softmax-on-margin pseudo-confidence with CalibratedClassifierCV probabilities plus a top-3 alternatives panel.")
bullet("Begin error analysis for the final write-up — confusion matrices for worst classes, residual conglomerate confusion pattern documented as a governance recommendation.")

doc.add_paragraph()
foot = doc.add_paragraph(); foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot.add_run("Submitted by Group 4, MGT 599 Capstone, DePaul University Chicago.")
r.font.size = Pt(9); r.italic = True; r.font.color.rgb = GRAY

doc.save(OUT)
print(f"Wrote {OUT}")
