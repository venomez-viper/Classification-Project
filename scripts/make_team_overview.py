from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)

RED   = RGBColor(0xC0, 0x39, 0x2B)
DARK  = RGBColor(0x1A, 0x1A, 0x2E)
GRAY  = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0x1E, 0x8B, 0x4C)
BLUE  = RGBColor(0x1A, 0x53, 0x9A)


def shd(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear")
    s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"), fill_hex)
    tcPr.append(s)


def para_shd(p, fill_hex):
    pPr = p._p.get_or_add_pPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear")
    s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"), fill_hex)
    pPr.append(s)


def h1(text, color=RED):
    p = doc.add_heading("", level=1)
    run = p.add_run(text)
    run.font.color.rgb = color
    run.font.size = Pt(14)
    run.font.bold = True
    return p


def h2(text, color=DARK):
    p = doc.add_heading("", level=2)
    run = p.add_run(text)
    run.font.color.rgb = color
    run.font.size = Pt(11.5)
    run.font.bold = True
    return p


def body(text, bold=False, color=None, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p


def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = DARK
    para_shd(p, "F2F2F2")
    return p


def table(headers, rows, hdr_color="C0392B", alt_color="FDF2F2"):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hr = t.rows[0]
    for i, h in enumerate(headers):
        c = hr.cells[i]
        c.text = h
        for run in c.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shd(c, hdr_color)
    for r, row_data in enumerate(rows):
        row = t.rows[r + 1]
        fill = alt_color if r % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            c = row.cells[ci]
            c.text = val
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(9.5)
            shd(c, fill)
    doc.add_paragraph()
    return t


def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("─" * 90)
    run.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)
    run.font.size = Pt(7)


# ══════════════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover.add_run("TAVSS")
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RED

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Taxonomy-Aware Venture Segmentation System")
run.font.size = Pt(14)
run.font.color.rgb = GRAY

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(
    "MGT 599 Capstone  ·  Group 4  ·  DePaul University Chicago  ·  Spring 2026"
)
run.font.size = Pt(11)
run.font.color.rgb = GRAY

doc.add_paragraph()

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run(
    "88.90% Macro F1  ·  145 Morningstar Classes  ·  Beats Fine-Tuned DeBERTa by +24.9 pp"
)
run.font.size = Pt(11)
run.font.bold = True
run.font.color.rgb = GREEN

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
h1("1. Project Overview")
body(
    "TAVSS automatically classifies company business descriptions into Morningstar Global Equity "
    "Classification Standard (GECS) codes. Given a plain-English description of what a company does, "
    "the system returns the correct 8-digit Morningstar industry code, a human-readable label, and a "
    "cross-mapped taxonomy (GICS, NAICS, SIC)."
)
doc.add_paragraph()

table(
    ["Item", "Detail"],
    [
        ["University",       "DePaul University Chicago"],
        ["Course",           "MGT 599 — Capstone Project"],
        ["Team",             "Group 4 — Akash, Srilaxmi, Vishal, Subasree, Tserennad"],
        ["Task 1 target",    "145 Morningstar industry codes (8-digit GECS)"],
        ["Task 2 target",    "407 Morningstar subindustry codes (10-digit GECS)"],
        ["Training data T1", "53,585 rows — SegmentName + SegmentDescription"],
        ["Training data T2", "22,012 rows — SegmentName + SegmentDescription"],
        ["Holdout test T1",  "10,717 rows — never seen during training"],
        ["Evaluation metric","Macro F1 (primary) + Weighted F1"],
        ["Rubric threshold", "75% Weighted F1 minimum to pass"],
    ],
    hdr_color="1A539A",
    alt_color="EEF4FB",
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. TECH STACK
# ══════════════════════════════════════════════════════════════════════════════
h1("2. Technology Stack")

h2("Machine Learning")
table(
    ["Library", "Version", "Used For"],
    [
        ["breezeml",       "0.2.5",   "Core ML framework — Akash's own published PyPI library"],
        ["scikit-learn",   "1.x",     "LinearSVC, TF-IDF vectoriser, metrics"],
        ["transformers",   "4.x",     "DeBERTa-v3-small tokeniser and model"],
        ["torch",          "2.x",     "GPU training loop for DeBERTa"],
        ["joblib",         "1.x",     "Model serialisation (.joblib files)"],
        ["scipy",          "1.x",     "Sparse CSR matrix operations"],
        ["pandas",         "2.x",     "Data loading and manipulation"],
        ["numpy",          "1.x",     "Numerical operations"],
    ],
    hdr_color="1E8B4C",
    alt_color="F0FAF4",
)

h2("Backend / API")
table(
    ["Component",        "Detail"],
    [
        ["Flask",        "Python microservice framework — 3 servers"],
        ["Waitress",     "WSGI production server — replaces Werkzeug on Windows"],
        ["Flask-CORS",   "Cross-origin requests from Next.js frontend"],
        ["server_legendary.py", "BreezeML Level 2 server — port 5003"],
        ["server_llm.py",       "DeBERTa inference server — port 5001"],
        ["server.py",           "Week 3 flat SVM baseline — port 5000"],
    ],
    hdr_color="7D3C98",
    alt_color="F9F0FF",
)

h2("Frontend")
table(
    ["Component",    "Detail"],
    [
        ["Next.js 14",      "React app-router framework"],
        ["TypeScript",      "Type-safe API calls and components"],
        ["Tailwind CSS",    "Utility-first styling"],
        ["Framer Motion",   "Animations and transitions"],
        ["Lucide React",    "Icon library"],
        ["npm run dev",     "Development server — port 3000"],
    ],
    hdr_color="C0392B",
    alt_color="FDF2F2",
)

# ══════════════════════════════════════════════════════════════════════════════
# 3. BREEZEML LIBRARY
# ══════════════════════════════════════════════════════════════════════════════
h1("3. breezeml — Our Own PyPI Library")
body(
    "breezeml (pip install breezeml) is a production-grade ML framework built on top of scikit-learn, "
    "authored by Akash Anipakalu Giridhar and published on PyPI. During this capstone we identified "
    "critical limitations and shipped 5 public patch versions to fix them."
)
doc.add_paragraph()

table(
    ["Version", "Fix", "Impact"],
    [
        ["v0.2.1", "joblib.Parallel(n_jobs=-1) — parallel benchmarking",           "O(N) → O(1) benchmark time"],
        ["v0.2.2", "X= and y= keyword args — bypass ColumnTransformer",            "Accepts scipy.sparse directly"],
        ["v0.2.3", "dual=False on all LinearSVC — primal formulation",             "20 min training → 2 seconds"],
        ["v0.2.4", "hasattr(model, 'save') — polymorphic save()",                  "Fixes fatal AttributeError"],
        ["v0.2.5", "class_weight='balanced' on all LinearSVC",                     "Macro F1: 43% → 86.82%"],
        ["Level 2","3-level cascade: Sector → Group → Code",                       "Macro F1: 59.70% → 88.90%"],
    ],
    hdr_color="C0392B",
    alt_color="FDF2F2",
)

code("from breezeml import classifiers")
code("model, report = classifiers.linear_svm(X=X_train, y=y_train, X_test=X_test, y_test=y_test)")

# ══════════════════════════════════════════════════════════════════════════════
# 4. WHAT WE BUILT — STEP BY STEP
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("4. What We Built — Step by Step")

h2("Week 1–2: Data Cleaning and Feature Engineering")
bullet("Loaded raw Morningstar GECS dataset — SegmentName, SegmentDescription, task codes")
bullet("Combined SegmentName + SegmentDescription into a single 'text' column (key insight: names carry direct class signal)")
bullet("Removed nulls, normalised code formats to 8-digit zero-padded strings")
bullet("Built descriptive analytics: class distribution, text length, rare class identification")
bullet("Identified severe class imbalance — some codes have 5,000 samples, others have < 10")
bullet("Output: data/cleaned/task1_clean.csv (53,585 rows) and task2_clean.csv (27,515 rows)")

doc.add_paragraph()
h2("Week 3: First Working Model — Flat SVM via breezeml")
bullet("TF-IDF vectoriser: 50,000 features, bigrams, sublinear_tf=True")
bullet("LinearSVC classifier via breezeml.classifiers.linear_svm()")
bullet("Patched breezeml v0.2.1 → v0.2.5 to fix performance bottlenecks (see Section 3)")
bullet("class_weight='balanced' was the critical fix — Macro F1 jumped from 43% to 86.82%")
bullet("80/20 stratified train/test split, random_state=42")
bullet("Serialised model to models/task1_svm_model.joblib")
bullet("Deployed Flask server on port 5000")

table(
    ["Metric", "Result"],
    [
        ["Weighted F1", "86.82%"],
        ["Macro F1",    "59.70%"],
        ["Accuracy",    "62.61%"],
        ["Rubric pass", "Yes — 86.82% > 75% threshold"],
    ],
    hdr_color="C0392B",
    alt_color="FDF2F2",
)

h2("Week 4: BreezeML Level 2 — Hierarchical Cascade")
bullet("Observed that flat SVM forces 'Oil & Gas Midstream' to compete against all 144 other codes")
bullet("Built taxonomy tree directly from Morningstar code structure (digits 1–3 = sector, 4–5 = group, 6–8 = code)")
bullet("Trained one LinearSVC per level: L1 (11 sectors), L2 (67 group models), L3 (145 code models)")
bullet("Shared single TF-IDF vectoriser across all three levels")
bullet("At inference time: L1 predicts sector → L2 predicts group within sector → L3 predicts final code")
bullet("Each classifier only competes within its branch (3–8 classes) instead of all 145")
bullet("Deployed BreezeML Level 2 server on port 5003")

table(
    ["Metric",                        "Flat SVM",  "BreezeML Level 2", "Change"],
    [
        ["Macro F1",                  "59.70%",    "88.90%",           "+29.2 pp"],
        ["Weighted F1",               "86.82%",    "88.90%",           "+2.1 pp"],
        ["Accuracy",                  "62.61%",    "89.11%",           "+26.5 pp"],
        ["Rare-class Macro F1 (≤10)", "20.44%",    "73.68%",           "+53.2 pp"],
        ["vs DeBERTa (64%)",          "−4.30 pp",  "+24.90 pp",        "—"],
        ["Speed",                     "~1,673/s",  "1,673/s",          "CPU only"],
    ],
    hdr_color="1E8B4C",
    alt_color="F0FAF4",
)

h2("Week 4: DeBERTa Fine-Tuning")
bullet("Fine-tuned microsoft/deberta-v3-small (180M parameters) on Task 1 (145 classes) and Task 2 (407 classes)")
bullet("Raw PyTorch training loop — no HuggingFace Trainer (more control, lower VRAM)")
bullet("Gradient accumulation (effective batch 16) to fit RTX 3050 4.3 GB VRAM")
bullet("Class-weighted CrossEntropyLoss to handle imbalance")
bullet("Adafactor optimiser on resume runs (saves ~500 MB VRAM vs AdamW)")
bullet("Task 1 result: 64.00% Macro F1 — beaten by cascade SVM by +24.9 pp")
bullet("Task 2 result: 32.49%+ Macro F1 (training ongoing at 12 epochs)")
bullet("Deployed DeBERTa server on port 5001")

# ══════════════════════════════════════════════════════════════════════════════
# 5. ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("5. System Architecture")

h2("BreezeML Level 2 — Cascade Inference Pipeline")
for line in [
    "Input: raw company description text",
    "         │",
    "         ▼  TF-IDF  (50,000 features, bigrams, sublinear_tf)",
    "         │",
    "         ▼  L1 — Broad Sector          ← single LinearSVC, 11 classes",
    '                  e.g.  "Energy"',
    "         │",
    "         ▼  L2 — Industry Group        ← one LinearSVC per sector (11 models)",
    '                  e.g.  "Oil & Gas"',
    "         │",
    "         ▼  L3 — Morningstar Code      ← one LinearSVC per group (67 models)",
    '                  e.g.  "Oil & Gas Exploration & Production"',
    "         │",
    "         ▼  Final 8-digit GECS code  +  analyst memo  +  GICS/NAICS/SIC crosswalk",
]:
    code(line)

doc.add_paragraph()
h2("Morningstar Code Structure")
for line in [
    "Code:  1  0  3  2  0  0  2  0",
    "       └──┬──┘  └──┬──┘  └──┬──┘",
    "          │         │         └── L3: Specific code  (digits 6–8)",
    "          │         └──────────── L2: Industry group (digits 4–5)",
    "          └────────────────────── L1: Broad sector   (digits 1–3)",
]:
    code(line)

doc.add_paragraph()
h2("DeBERTa Inference Pipeline")
for line in [
    "Input: plain-English company description (no jargon needed)",
    "         │",
    "         ▼  SentencePiece tokenisation  (128K vocabulary)",
    "         │",
    "         ▼  12 attention layers  (cross-word meaning, not keywords)",
    "         │",
    "         ▼  Classification head  (softmax over 145 or 407 classes)",
    "         │",
    "         ▼  Top-3 predictions + confidence scores",
]:
    code(line)

doc.add_paragraph()
h2("Deployed Servers")
table(
    ["Server file",         "Port", "Model",                "Start command"],
    [
        ["server_legendary.py", "5003", "BreezeML Level 2 cascade", "python server_legendary.py"],
        ["server_llm.py",       "5001", "DeBERTa-v3-small",         "python server_llm.py"],
        ["server.py",           "5000", "Flat SVM baseline (Week 3)","python server.py"],
        ["frontend/",           "3000", "Next.js web interface",     "npm run dev"],
    ],
    hdr_color="7D3C98",
    alt_color="F9F0FF",
)

# ══════════════════════════════════════════════════════════════════════════════
# 6. SVM VS DEBERTA
# ══════════════════════════════════════════════════════════════════════════════
h1("6. SVM vs DeBERTa — Architecture Beat Intelligence")

body(
    "DeBERTa is the more intelligent model. It reads meaning — 'workers drill deep holes into the earth' "
    "correctly maps to Oil & Gas Exploration even with zero financial jargon. The SVM cannot do this. "
    "Yet the SVM won by +24.9 percentage points. Why?"
)
doc.add_paragraph()

h2("The structural advantage")
body(
    "A flat model forces every class to compete against all 144 others in a single decision. "
    "BreezeML Level 2 mirrors the taxonomy tree — at each level only 3–8 closely related classes "
    "compete. Rare classes no longer lose to dominant ones. This one structural change "
    "added +53 pp on rare-class F1 alone."
)
doc.add_paragraph()

table(
    ["Capability",                  "BreezeML Level 2",         "DeBERTa"],
    [
        ["Understands plain English",   "No — needs keywords",      "Yes — reads meaning"],
        ["Taxonomy-aware structure",    "Yes — 3-level hierarchy",  "No — flat softmax"],
        ["Task 1 Macro F1",             "88.90%",                   "64.00%"],
        ["Rare-class F1",               "73.68%",                   "~20% (estimated)"],
        ["Speed",                       "1,673 samples/sec CPU",    "~40/sec GPU"],
        ["GPU required",                "No",                       "Yes"],
    ],
    hdr_color="1A539A",
    alt_color="EEF4FB",
)

body(
    "Key finding: a well-structured simpler model outperforms a powerful model given the wrong "
    "structure — especially on imbalanced, hierarchical classification problems.",
    bold=True,
    color=RED,
)

# ══════════════════════════════════════════════════════════════════════════════
# 7. FILE STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("7. Project File Structure")

for line in [
    "capstone MGT 599/",
    "│",
    "├── server.py                   ← Week 3 flat SVM (port 5000)",
    "├── server_llm.py               ← DeBERTa microservice (port 5001)",
    "├── server_legendary.py         ← BreezeML Level 2 full stack (port 5003)",
    "│",
    "├── scripts/",
    "│   ├── train_cascade.py        ← train BreezeML Level 2",
    "│   ├── cascade_predict.py      ← inference engine",
    "│   ├── cascade_common.py       ← taxonomy tree builder",
    "│   └── benchmark.py            ← compare all models",
    "│",
    "├── models/",
    "│   ├── cascade_L1_svm.joblib   ← sector classifier (11 classes)",
    "│   ├── cascade_L2_models.joblib← industry group classifiers (11 models)",
    "│   ├── cascade_L3_models.joblib← code classifiers (67 models)",
    "│   ├── cascade_vectorizer.pkl  ← shared TF-IDF vectoriser",
    "│   ├── task1_svm_model.joblib  ← Week 3 flat SVM",
    "│   └── task1_tfidf_vectorizer.pkl",
    "│",
    "├── legendary/",
    "│   ├── shared.py               ← label lookup",
    "│   ├── inference_router.py     ← confidence-based routing logic",
    "│   ├── explanations.py         ← heuristic analyst memo generator",
    "│   └── taxonomy_crosswalk.py   ← GICS / NAICS / SIC mapping",
    "│",
    "├── llm_finetuning/",
    "│   ├── scripts/train_local.py  ← DeBERTa training script",
    "│   ├── data/task1_train.csv    ← 42,868 training rows",
    "│   ├── data/task2_train.csv    ← 22,012 training rows",
    "│   └── results/",
    "│       ├── task1_best_model/   ← DeBERTa Task 1 checkpoint",
    "│       └── task2_best_model/   ← DeBERTa Task 2 checkpoint",
    "│",
    "├── data/cleaned/",
    "│   ├── task1_clean.csv         ← 53,585 rows, 145 classes",
    "│   └── task2_clean.csv         ← 27,515 rows, 407 classes",
    "│",
    "├── frontend/                   ← Next.js web application",
    "│   ├── app/demo/               ← BreezeML Level 2 demo page",
    "│   ├── app/llm/                ← DeBERTa demo page",
    "│   └── app/dashboard/          ← ML metrics dashboard",
    "│",
    "└── docs/",
    "    ├── model_version_history.md",
    "    ├── Week3_Model_Architecture_and_Pipeline.md",
    "    ├── Week4_BreezeML_Level2_Team_Guide.md",
    "    ├── Week4_Team_Tasks.md",
    "    └── SVM_vs_DeBERTa_Analysis.md",
]:
    code(line)

# ══════════════════════════════════════════════════════════════════════════════
# 8. HOW TO RUN EVERYTHING
# ══════════════════════════════════════════════════════════════════════════════
h1("8. How to Run Everything")

h2("Train BreezeML Level 2")
code('cd "C:\\Users\\akash\\Desktop\\capstone MGT 599"')
code("python scripts/train_cascade.py")
body("Trains in 2–5 minutes on CPU. Writes 5 model files to models/.")

doc.add_paragraph()
h2("Run the Benchmark")
code("python scripts/benchmark.py")
body("Compares Level 2 vs flat SVM vs DeBERTa on the 10,717-sample holdout test set.")

doc.add_paragraph()
h2("Start All Servers")
code("python server_legendary.py      # BreezeML Level 2  — port 5003")
code("python server_llm.py            # DeBERTa           — port 5001")
code("python server.py                # Flat SVM baseline — port 5000")
code("cd frontend && npm run dev      # Web interface     — port 3000")

doc.add_paragraph()
h2("Test BreezeML Level 2 via PowerShell")
code('$b = \'{"text": "The company provides retail banking and mortgage loans."}\'')
code("Invoke-WebRequest -Uri http://localhost:5003/api/predict_legendary -Method POST -Body $b -ContentType 'application/json' | Select-Object -ExpandProperty Content")

doc.add_paragraph()
h2("Train DeBERTa Task 2 (from scratch)")
code("python llm_finetuning/scripts/train_local.py --task task2")
body("~3–4 hours on RTX 3050. Saves best checkpoint to llm_finetuning/results/task2_best_model/")

doc.add_paragraph()
h2("Continue DeBERTa Training (resume from checkpoint)")
code("python llm_finetuning/scripts/train_local.py --task task2 --resume")

# ══════════════════════════════════════════════════════════════════════════════
# 9. FINAL RESULTS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("9. Final Results Summary")

table(
    ["Model", "Macro F1", "Weighted F1", "Accuracy", "Notes"],
    [
        ["BreezeML Level 2 ★", "88.90%", "88.90%", "89.11%", "Champion — CPU only"],
        ["DeBERTa-v3-small",   "64.00%", "—",      "—",      "Task 1 fine-tuned"],
        ["Flat SVM (Week 3)",  "59.70%", "86.82%", "62.61%", "breezeml v0.2.5"],
        ["Random baseline",    "0.69%",  "—",      "—",      "1/145 classes"],
    ],
    hdr_color="1E8B4C",
    alt_color="F0FAF4",
)

h2("Key Numbers to Remember")
table(
    ["Stat",                          "Value"],
    [
        ["Task 1 classes",            "145 Morningstar GECS codes"],
        ["Task 2 classes",            "407 subindustry codes"],
        ["Holdout test samples",      "10,717 (Task 1)"],
        ["Cascade vs flat gain",      "+29.2 pp Macro F1"],
        ["Cascade vs DeBERTa gain",   "+24.9 pp Macro F1"],
        ["Rare-class gain",           "+53.2 pp (20.44% → 73.68%)"],
        ["Throughput",                "1,673 classifications per second on CPU"],
        ["breezeml versions shipped", "5 public PyPI releases (v0.2.1 → v0.2.5)"],
        ["Rubric threshold",          "75% Weighted F1 — achieved 88.90%"],
    ],
    hdr_color="C0392B",
    alt_color="FDF2F2",
)

# ══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
fp = doc.add_paragraph(
    "MGT 599 Capstone  ·  Group 4  ·  DePaul University Chicago  ·  Spring 2026  ·  BreezeML Level 2"
)
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.runs[0].font.size = Pt(9)
fp.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

out = r"C:\Users\akash\Desktop\capstone MGT 599\docs\TAVSS_Project_Overview.docx"
doc.save(out)
print("Saved:", out)
