from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)


def shaded_para(p, fill_hex):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def code_line(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text if text.strip() else " ")
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    shaded_para(p, "F4F4F4")


def code_block(src):
    for line in src.split("\n"):
        code_line(line)
    doc.add_paragraph()


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "C0392B")
        tcPr.append(shd)
    for r, row_data in enumerate(rows):
        row = table.rows[r + 1]
        for c, val in enumerate(row_data):
            row.cells[c].text = val
    doc.add_paragraph()
    return table


# ── Title ────────────────────────────────────────────────────────────────────
t = doc.add_heading("Week 4 — Team Task Sheet", 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.runs[0].font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

sub = doc.add_paragraph("MGT 599 Capstone  ·  Group 4  ·  DePaul University Chicago")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(12)
sub.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
doc.add_paragraph()

# Intro
p = doc.add_paragraph()
p.add_run("Each person runs ").font.size = Pt(11)
r = p.add_run("one independent model")
r.bold = True
r.font.size = Pt(11)
p.add_run(" this week. No coordination needed — run your script, record your numbers, send to Akash.").font.size = Pt(11)

p2 = doc.add_paragraph()
r2 = p2.add_run("Do NOT touch: ")
r2.bold = True
r2.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
p2.add_run("server_legendary.py, train_cascade.py, or anything in legendary/ — that is Akash's work.")
doc.add_paragraph()

# ── Assignment table ─────────────────────────────────────────────────────────
doc.add_heading("Who Does What", level=1)
add_table(
    ["Person", "Model", "Task", "Script to run"],
    [
        ["Srilaxmi",  "Linear SVM",           "Task 1 — 145 classes", "python scripts/srilaxmi_week4.py"],
        ["Vishal",    "Logistic Regression",   "Task 1 — 145 classes", "python scripts/vishal_week4.py"],
        ["Subasree",  "Linear SVM (balanced)", "Task 2 — 407 classes", "python scripts/subasree_week4.py"],
        ["Tserennad", "Naive Bayes",           "Task 1 — 145 classes", "python scripts/tserennad_week4.py"],
    ],
)

# ── Setup ────────────────────────────────────────────────────────────────────
doc.add_heading("Setup — Everyone Does This First", level=1)
code_block(
    'cd "C:\\Users\\akash\\Desktop\\capstone MGT 599"\n'
    "pip install breezeml==0.2.5\n"
    'python -c "import breezeml; print(breezeml.__version__)"  # should print 0.2.5'
)

SCORE_HEADERS = ["Metric", "Your Result", "Week 3 Baseline"]

# ── Srilaxmi ─────────────────────────────────────────────────────────────────
doc.add_heading("Srilaxmi — Linear SVM on Task 1", level=1)
doc.add_paragraph("Model: breezeml LinearSVC  |  Data: task1_clean.csv  |  Classes: 145  |  Beat: 59.70% Macro F1")
code_block(
    "# save as: scripts/srilaxmi_week4.py\n"
    "# run:     python scripts/srilaxmi_week4.py\n"
    "\n"
    "import pandas as pd\n"
    "from sklearn.feature_extraction.text import TfidfVectorizer\n"
    "from sklearn.model_selection import train_test_split\n"
    "from breezeml import classifiers\n"
    "\n"
    'df  = pd.read_csv("data/cleaned/task1_clean.csv")\n'
    'df  = df.dropna(subset=["text", "mstar_code"])\n'
    "\n"
    "vec = TfidfVectorizer(max_features=50_000, ngram_range=(1, 2), sublinear_tf=True)\n"
    'X   = vec.fit_transform(df["text"].tolist())\n'
    'y   = df["mstar_code"].astype(str).tolist()\n'
    "\n"
    "X_train, X_test, y_train, y_test = train_test_split(\n"
    "    X, y, test_size=0.2, stratify=y, random_state=42\n"
    ")\n"
    "\n"
    "_, report = classifiers.linear_svm(X=X_train, y=y_train, X_test=X_test, y_test=y_test)\n"
    "print(report)"
)
doc.add_paragraph("Record your score:").runs[0].bold = True
add_table(SCORE_HEADERS, [["Macro F1", "______", "59.70%"], ["Weighted F1", "______", "86.82%"]])

# ── Vishal ───────────────────────────────────────────────────────────────────
doc.add_heading("Vishal — Logistic Regression on Task 1", level=1)
doc.add_paragraph("Model: breezeml Logistic Regression  |  Data: task1_clean.csv  |  Classes: 145  |  Beat: 59.70% Macro F1")
code_block(
    "# save as: scripts/vishal_week4.py\n"
    "# run:     python scripts/vishal_week4.py\n"
    "\n"
    "import pandas as pd\n"
    "from sklearn.feature_extraction.text import TfidfVectorizer\n"
    "from sklearn.model_selection import train_test_split\n"
    "from breezeml import classifiers\n"
    "\n"
    'df  = pd.read_csv("data/cleaned/task1_clean.csv")\n'
    'df  = df.dropna(subset=["text", "mstar_code"])\n'
    "\n"
    "vec = TfidfVectorizer(max_features=50_000, ngram_range=(1, 2), sublinear_tf=True)\n"
    'X   = vec.fit_transform(df["text"].tolist())\n'
    'y   = df["mstar_code"].astype(str).tolist()\n'
    "\n"
    "X_train, X_test, y_train, y_test = train_test_split(\n"
    "    X, y, test_size=0.2, stratify=y, random_state=42\n"
    ")\n"
    "\n"
    "_, report = classifiers.logistic_regression(X=X_train, y=y_train, X_test=X_test, y_test=y_test)\n"
    "print(report)"
)
doc.add_paragraph("Record your score:").runs[0].bold = True
add_table(SCORE_HEADERS, [["Macro F1", "______", "59.70%"], ["Weighted F1", "______", "86.82%"]])

# ── Subasree ─────────────────────────────────────────────────────────────────
doc.add_heading("Subasree — Linear SVM on Task 2", level=1)
doc.add_paragraph("Model: breezeml LinearSVC (balanced)  |  Data: task2_clean.csv  |  Classes: 407  |  Beat: 47.72% Weighted F1")
code_block(
    "# save as: scripts/subasree_week4.py\n"
    "# run:     python scripts/subasree_week4.py\n"
    "\n"
    "import pandas as pd\n"
    "from sklearn.feature_extraction.text import TfidfVectorizer\n"
    "from sklearn.model_selection import train_test_split\n"
    "from breezeml import classifiers\n"
    "\n"
    'df  = pd.read_csv("data/cleaned/task2_clean.csv")\n'
    'df  = df.dropna(subset=["text"])\n'
    "\n"
    'label_col = "GECSSubIndustryCode"   # update if your column name differs\n'
    "y   = df[label_col].astype(str).tolist()\n"
    "\n"
    "vec = TfidfVectorizer(max_features=10_000, ngram_range=(1, 2), sublinear_tf=True)\n"
    'X   = vec.fit_transform(df["text"].tolist())\n'
    "\n"
    "X_train, X_test, y_train, y_test = train_test_split(\n"
    "    X, y, test_size=0.2, stratify=y, random_state=42\n"
    ")\n"
    "\n"
    "_, report = classifiers.linear_svm(\n"
    "    X=X_train, y=y_train, X_test=X_test, y_test=y_test,\n"
    '    class_weight="balanced"\n'
    ")\n"
    "print(report)"
)
doc.add_paragraph("Record your score:").runs[0].bold = True
add_table(SCORE_HEADERS, [["Macro F1", "______", "39.62%"], ["Weighted F1", "______", "47.72%"]])

# ── Tserennad ────────────────────────────────────────────────────────────────
doc.add_heading("Tserennad — Naive Bayes on Task 1", level=1)
doc.add_paragraph("Model: breezeml Naive Bayes  |  Data: task1_clean.csv  |  Classes: 145  |  Beat: 59.70% Macro F1")
code_block(
    "# save as: scripts/tserennad_week4.py\n"
    "# run:     python scripts/tserennad_week4.py\n"
    "\n"
    "import pandas as pd\n"
    "from sklearn.feature_extraction.text import TfidfVectorizer\n"
    "from sklearn.model_selection import train_test_split\n"
    "from breezeml import classifiers\n"
    "\n"
    'df  = pd.read_csv("data/cleaned/task1_clean.csv")\n'
    'df  = df.dropna(subset=["text", "mstar_code"])\n'
    "\n"
    "vec = TfidfVectorizer(max_features=50_000, ngram_range=(1, 2), sublinear_tf=True)\n"
    'X   = vec.fit_transform(df["text"].tolist())\n'
    'y   = df["mstar_code"].astype(str).tolist()\n'
    "\n"
    "X_train, X_test, y_train, y_test = train_test_split(\n"
    "    X, y, test_size=0.2, stratify=y, random_state=42\n"
    ")\n"
    "\n"
    "_, report = classifiers.naive_bayes(X=X_train, y=y_train, X_test=X_test, y_test=y_test)\n"
    "print(report)"
)
doc.add_paragraph("Record your score:").runs[0].bold = True
add_table(SCORE_HEADERS, [["Macro F1", "______", "59.70%"], ["Weighted F1", "______", "86.82%"]])

# ── Combined results ─────────────────────────────────────────────────────────
doc.add_heading("Combined Results — Fill In and Send to Akash", level=1)
add_table(
    ["Person", "Model", "Task", "Macro F1", "Weighted F1"],
    [
        ["Srilaxmi",          "Linear SVM",           "Task 1", "______", "______"],
        ["Vishal",            "Logistic Regression",   "Task 1", "______", "______"],
        ["Subasree",          "Linear SVM (balanced)", "Task 2", "______", "______"],
        ["Tserennad",         "Naive Bayes",           "Task 1", "______", "______"],
        ["Baseline (Week 3)", "Flat SVM",              "Task 1", "59.70%", "86.82%"],
    ],
)

# Footer
fp = doc.add_paragraph("MGT 599 Capstone  ·  Group 4  ·  DePaul University Chicago  ·  Week 4  ·  May 2026")
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.runs[0].font.size = Pt(9)
fp.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

out = r"C:\Users\akash\Desktop\capstone MGT 599\docs\Week4_Team_Tasks.docx"
doc.save(out)
print("Saved:", out)
