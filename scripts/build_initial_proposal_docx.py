"""Build a polished Word doc of the Initial Proposal."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
OUT  = ROOT / "docs" / "Initial_Proposal.docx"

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

# ── Default font ──
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
TEAL = RGBColor(0x0E, 0x6B, 0x6E)


def heading(text, level=1, color=NAVY):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        if level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(14)
    return h


def para(text, bold=False, italic=False, size=11, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def hrule():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)


def shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def styled_table(rows, header=True, fill_header="1F3A5F", fill_header_text="FFFFFF",
                 col_widths=None):
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            if r_idx == 0 and header:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                shade_cell(cell, fill_header)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for c_idx, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[c_idx].width = w
    return table


# ─────────────────────────────────────────────────────────────────────────────
# COVER
# ─────────────────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Initial Project Proposal")
run.font.size = Pt(28)
run.font.color.rgb = NAVY
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("GECS Industry and Business Activity Classification\nusing Machine Learning")
run.font.size = Pt(16)
run.font.color.rgb = TEAL
run.italic = True

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line, bold in [
    ("MGT 599 Capstone Project · Q2 2026", True),
    ("Group 4 · DePaul University Chicago", False),
    ("Lead: Akash Anipakalu Giridhar", False),
    ("Team: Srilaxmi, Vishal, Subasree, Tserennad", False),
    ("Industry Partner: Morningstar — Reference Entity Data (RED) Team", False),
]:
    r = meta.add_run(line + "\n")
    r.font.size = Pt(12)
    r.bold = bold

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = date_p.add_run("Date: May 10, 2026")
r.font.size = Pt(11)
r.italic = True

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 1. EXECUTIVE SUMMARY
# ─────────────────────────────────────────────────────────────────────────────
heading("1. Executive Summary", 1)
para("This proposal documents our first five weeks of work on the Morningstar GECS classification capstone, "
     "presents our preliminary findings and the substantial methodological audit we completed, and lays out "
     "the technical path forward through the remaining sprint.")
para("We are building two supervised classifiers:")
bullet("Task 1 — GECS Industry Classification: map a company's LongProfile + segment text into one of 145 Morningstar industry codes.")
bullet("Task 2 — Business Activity Subindustry Classification: map segment-level text into one of 450 business activity codes, exploiting the deterministic one-to-many parent–child relationship from Task 1 to Task 2.")
para("Our baseline pipeline reached an apparent 88.90% Macro F1 in Week 3. In Week 4 we discovered this number "
     "was inflated by training/test leakage — 97.2% of test rows were present in the training set. The honest "
     "baseline, after rebuilding the pipeline with strict train/test separation, is 69.09% Macro F1.")
para("This proposal explains exactly how we got here, what we have learned, why we are not aiming for the "
     "illusion of 88.90% anymore, and what we will do in the remaining weeks to push the honest number toward "
     "the 75–80% range expected by the case rubric.")
hrule()

# ─────────────────────────────────────────────────────────────────────────────
# 2. BACKGROUND
# ─────────────────────────────────────────────────────────────────────────────
heading("2. Problem Background", 1)
para("Morningstar's Reference Entity Data (RED) team maintains GECS, a four-level hierarchical taxonomy that "
     "organizes every publicly listed company in their global coverage universe into a sector → industry "
     "group → industry → business activity tree. Analysts use GECS for peer comparison, portfolio construction, "
     "risk attribution, and reporting. Misclassification at the leaf level propagates upward and distorts every "
     "product downstream — including PitchBook on the private-markets side.")
para("GECS classification today is largely manual. The case asks us whether modern NLP can scale this work: "
     "given the structured and unstructured company disclosures Morningstar already collects, can we automate "
     "the assignment with production-grade accuracy?")
para("The case is concrete:")
bullet("3 Super Sectors → 11 Sectors → 55 Industry Groups → 145 Industries → 450 Business Activities")
bullet("Task 1 dataset: 53,585 records, Dec 2003 – Dec 2024, 145 classes")
bullet("Task 2 dataset: 27,537 records, May 2020 – Dec 2024, 450 classes")
bullet("Stated pass-criteria: Macro F1 ≥ 0.75 on overall classification; top-10 most-frequent class F1 > 0.85")
hrule()

# ─────────────────────────────────────────────────────────────────────────────
# 3. OBJECTIVES
# ─────────────────────────────────────────────────────────────────────────────
heading("3. Objectives", 1)
styled_table([
    ["#", "Objective", "Status"],
    ["1", "Stand up an honest end-to-end training and evaluation pipeline for both tasks", "Complete (Week 4)"],
    ["2", "Document dataset structure, label distributions, and class imbalance", "Complete (Week 2)"],
    ["3", "Produce reproducible baselines using classical ML (TF-IDF + linear classifiers)", "Complete (Week 3)"],
    ["4", "Audit baselines for hidden data leakage and methodological errors", "Complete (Week 4)"],
    ["5", "Move beyond bag-of-words representations with semantic embeddings", "Complete (Week 4)"],
    ["6", "Hit Macro F1 ≥ 0.75 on Task 1", "In progress (Week 5)"],
    ["7", "Build Task 2 classifier exploiting the Task 1 → Task 2 constraint", "Scaffolding underway (Week 5)"],
    ["8", "Deliver demo, full audit, and reproducible code", "Ongoing"],
])
hrule()

# ─────────────────────────────────────────────────────────────────────────────
# 4. DATA OVERVIEW
# ─────────────────────────────────────────────────────────────────────────────
heading("4. Data Overview", 1)

heading("4.1 Task 1 — Industry classification", 2)
styled_table([
    ["Field", "Type", "Used as"],
    ["CompanyId", "string", "Group key"],
    ["AsOfDate", "string", "Snapshot stamp"],
    ["LongProfile", "free text", "Company-level description"],
    ["SegmentName", "string", "Segment label"],
    ["SegmentDescription", "free text", "Segment-level description"],
    ["Revenue", "float", "Engineered feature input"],
    ["total_revenue_company_as_of", "float", "Engineered feature input"],
    ["revenue_share", "float", "Engineered feature"],
    ["is_largest_share_segment", "bool", "Engineered feature"],
    ["MstarGlobal", "class (145)", "Target"],
])

heading("4.2 Key data characteristics", 2)
bullet("Class imbalance: top-10 codes account for ~30% of records; bottom-50 codes have fewer than 200 records each.")
bullet("Multi-segment companies are common: ~35% of CompanyId values have segments mapping to different MstarGlobal codes. These are diversified conglomerates whose LongProfile describes the entire company but whose individual segments resolve to different industries.")
bullet("Sector 310 (Industrials) contains the single hardest leaf code, 31030010 (Diversified Industrial Conglomerates), F1 ~ 30% on honest evaluation.")

heading("4.3 Task 2", 2)
para("Smaller dataset (27.5k rows), more classes (450), shorter text (SegmentName + SegmentDescription only — "
     "no LongProfile). Each Task 2 code deterministically rolls up into exactly one Task 1 code. This "
     "relationship is a hard constraint we will use at inference time.")
hrule()

# ─────────────────────────────────────────────────────────────────────────────
# 5. CODE JOURNEY
# ─────────────────────────────────────────────────────────────────────────────
heading("5. The Code Journey — Weeks 1 through 5", 1)
para("This section is the spine of the proposal: the work, in order, and what each iteration taught us.")

# Week 1
heading("Week 1 — Project setup", 2)
bullet("Team formation, environment standardization (Python 3.11, sklearn 1.4, custom breezeml wrapper).")
bullet("Repo layout: data/, scripts/, models/, notebooks/, docs/.")
bullet("Initial reading of the case PDF and the Morningstar GECS taxonomy reference.")

# Week 2
heading("Week 2 — Data exploration and cleaning", 2)
bullet("Built task1_clean.csv and task2_clean.csv.")
bullet("Profiled class distributions, text length distributions, missing-value patterns.")
bullet("Engineered initial structured features: revenue_share, is_largest_share_segment, plus num_segments, max_share, share_std at the company level.")
bullet("Deliverable: docs/Week2_What_We_Did.md + descriptive-analytics notebook.")

# Week 3
heading("Week 3 — Baseline classical-ML pipeline", 2)
bullet("Implemented the canonical TF-IDF + Linear SVM pipeline via our breezeml library wrapper.")
bullet("Task 1: Weighted F1 = 86.82%, Macro F1 = 61.07% (random 80/20 stratified split).")
bullet("Task 2: Weighted F1 = 47.72%, Macro F1 = 39.62% (407 classes after filtering single-occurrence labels).")
bullet("Found interpretable industry-defining vocabulary (semiconductor, brokerage, pharmaceutical). Misclassifications were near-misses inside the correct sector.")
bullet("Deliverable: docs/Week3_Report.md, week3_modeling_task1.ipynb, week3_modeling_task2.ipynb.")

# Week 4
heading("Week 4 — The audit that changed everything", 2)
para("We attempted to ship a 'legendary' cascaded SVM (sector → group → industry) and an interactive demo. "
     "The cascade reported 88.90% Macro F1 on Task 1 — an apparent jump of nearly 30 points over the Week 3 "
     "baseline. The demo, however, only produced sensible predictions on four hand-crafted example inputs. "
     "Arbitrary user input returned random-looking labels with fake high 'confidence' numbers.")
para("That mismatch forced an audit. What we found:", bold=True)
bullet("Training/test leakage in the original cascade. The model trained on data/cleaned/task1_clean.csv (53,585 rows = the full dataset) and was evaluated on llm_finetuning/data/task1_test.csv (10,717 rows). 97.2% of those test rows were present in training; only 305 were truly unseen. On the unseen 305, the same model scored 81.73%, not 88.90%. The model is real — but the headline metric was leaked memorization.")
bullet("Fake confidence display. The demo rendered softmax(SVM decision-function margin) as a 'confidence' percentage. For out-of-distribution input the SVM still produces decision margins, the softmax still normalizes them, and the UI happily reported '92% confident' while predicting wrongly.")
bullet("Conglomerate noise. ~35% of training companies are multi-segment with multiple labels. Because we concatenated LongProfile with segment text, an identical LongProfile prefix appears in ~55% of training rows mapping to different MstarGlobal codes. That is irreducible label noise we manufactured in preprocessing.")
bullet("Cascade error propagation. L1 (sector) errors cascade downward. ~52% of all final errors trace back to a wrong L1 prediction. A top-down cascade is structurally bad for this problem unless L1 is near-perfect.")
bullet("TF-IDF ceiling. Even after fixing leakage, pure TF-IDF + LinearSVC plateaus at ~57% Macro F1, regardless of vocabulary size, char n-grams, or C tuning.")
para("The full audit is preserved in CASCADE_AUDIT.md.")
para("We responded by rebuilding the pipeline honestly, then explored three more ambitious directions:", bold=True)
bullet("V13: GECS official-taxonomy anchoring. Parsed all 145 industry definitions from the Morningstar 2019 GECS PDF, encoded each definition with MiniLM and BGE, and added cosine similarity to every official anchor as 580 extra features. Result: 67.99% Macro F1.")
bullet("V14: Retrieval-Augmented Classification (RAC). Top-25 nearest training rows aggregated into a 145-class prior. Result: 66.04%.")
bullet("V16: FinBERT fine-tuning. Three epochs on Google Colab T4 GPU using yiyanghkust/finbert-pretrain. Result: 61.84% — domain mismatch (FinBERT trained on financial news, not company descriptions).")
para("Deliverables: CASCADE_AUDIT.md, PROJECT_JOURNEY.md, 17 numbered training scripts, all reproducible.")

# Week 5
heading("Week 5 — The hypothesis we are testing now", 2)
para("The audit produced one diagnostic insight that we believe explains the universal ~68% plateau: for "
     "~55% of training rows, the input text is contaminated with the same LongProfile prefix as several "
     "other rows that have different labels. No encoder, no loss function, and no ensemble can recover the "
     "right answer from input that ambiguously points to multiple labels.")
para("Our Week 5 plan moves in two directions simultaneously:", bold=True)
bullet("Validate the contamination hypothesis by training every Week 5 model on segment text only (SegmentName + SegmentDescription), with LongProfile either dropped or used only as a low-weighted auxiliary signal.")
bullet("Apply hierarchy-aware modeling on top of clean inputs: a single transformer encoder (DeBERTa-v3-base) with multi-task heads for sector (11), group (55), and industry (145), trained with a joint hierarchy-weighted loss. Add Distribution-Balanced loss to lift rare-class macro F1.")

para("Team distribution this week (see docs/Week5_Team_Tasks.md for executable scripts):", bold=True)
styled_table([
    ["Member", "Model", "Task", "Purpose"],
    ["Srilaxmi", "Linear SVM with word + character n-grams", "Task 1", "Test if subword signal helps rare codes"],
    ["Vishal", "Logistic Regression with 100k vocab + trigrams", "Task 1", "Test if higher-capacity linear model helps"],
    ["Subasree", "Linear SVM with class_weight='balanced'", "Task 2", "Lift rare-subindustry F1"],
    ["Tserennad", "Random Forest", "Task 1", "Non-linear baseline"],
    ["Akash (lead)", "DeBERTa-v3 hierarchical multi-task + segment-only", "Task 1 & 2", "Strategic core experiment"],
])
hrule()

# ─────────────────────────────────────────────────────────────────────────────
# 6. RESULTS
# ─────────────────────────────────────────────────────────────────────────────
heading("6. Preliminary Results to Date", 1)
styled_table([
    ["Pipeline", "Split", "Macro F1", "Accuracy", "Top-10"],
    ["Original cascade (Week 3)", "Leaked (train ⊇ test)", "88.90%", "n/a", "9/10"],
    ["Original cascade on truly unseen rows", "Honest (n=305)", "81.73%", "n/a", "n/a"],
    ["Honest TF-IDF cascade (Week 4 rebuild)", "task1_train → test", "59.65%", "62.0%", "1/10"],
    ["TF-IDF + numerical engineered", "task1_train → test", "63.42%", "66.2%", "1/10"],
    ["MiniLM embeddings only", "task1_train → test", "59.70%", "62.0%", "1/10"],
    ["V5 hybrid (TF-IDF + MiniLM + numerical)", "task1_train → test", "67.11%", "70.1%", "2/10"],
    ["V6 hybrid (TF-IDF + BGE + numerical)", "task1_train → test", "67.70%", "70.5%", "2/10"],
    ["V8 mega-ensemble", "task1_train → test", "68.42%", "~71%", "2/10"],
    ["V10 calibrated stack (current best)", "task1_train → test", "69.09%", "71.65%", "2/10"],
    ["V13 (V8 + GECS PDF anchors)", "task1_train → test", "67.99%", "70.81%", "2/10"],
    ["V14 (Retrieval-Augmented)", "task1_train → test", "66.04%", "68.69%", "1/10"],
    ["V16 FinBERT 3-epoch (Colab)", "task1_train → test", "61.84%", "62.15%", "0/10"],
])
hrule()

# ─────────────────────────────────────────────────────────────────────────────
# 7. FINDINGS
# ─────────────────────────────────────────────────────────────────────────────
heading("7. Findings to Date — What We Have Actually Learned", 1)
bullet("Naïve cascade evaluation is dangerous. Anyone reporting 88%+ on this task without a rigorous train/test audit is almost certainly leaking. We caught ourselves doing this.")
bullet("The representation isn't the bottleneck. TF-IDF (60%), MiniLM (60%), and BGE-base (60%) all plateau in the same place. Switching encoders gets diminishing returns.")
bullet("Engineered features matter. Adding num_segments, max_share, share_std on top of the TF-IDF + embedding stack added +4 to +7 percentage points by itself.")
bullet("The data has manufactured label noise. Concatenating LongProfile with segment text creates many-to-many mapping for conglomerate companies. This is the most plausible explanation for the universal ~68% ceiling across architectures.")
bullet("The official GECS taxonomy is unused signal. No team will think to parse the Morningstar 2019 GECS PDF and use it as semantic anchors. We have implemented this; we expect it to contribute meaningfully once the input contamination is fixed.")
bullet("Domain-pretrained BERT is not automatically better. FinBERT was pretrained on financial news, not company descriptions. The distribution mismatch hurt us. The encoder choice has to fit the text type.")
hrule()

# ─────────────────────────────────────────────────────────────────────────────
# 8. INNOVATIONS
# ─────────────────────────────────────────────────────────────────────────────
heading("8. Methodological Innovations", 1)
para("The parts of this work that are genuinely novel for this dataset:")
bullet("End-to-end leakage audit. We caught and documented a 30-percentage-point inflation in our own baseline. The audit document (CASCADE_AUDIT.md) reads like a postmortem and is itself a deliverable.")
bullet("Official taxonomy anchoring. Parsing the Morningstar 2019 GECS structure PDF and using the 145 official industry definitions as semantic anchors is not standard practice for this dataset.")
bullet("Multi-encoder hybrid feature stack with empirical class prototypes. Stacking TF-IDF + MiniLM + BGE + class centroids + numerical features into a single ~123k-dimensional sparse representation, calibrated through CalibratedClassifierCV, has produced our current best honest result.")
bullet("Honest probability display. We are replacing the original demo's softmax-on-margin pseudo-confidence with calibrated probabilities, plus a top-3 alternatives panel for low-confidence predictions.")
hrule()

# ─────────────────────────────────────────────────────────────────────────────
# 9. RISKS
# ─────────────────────────────────────────────────────────────────────────────
heading("9. Risks and Mitigations", 1)
styled_table([
    ["Risk", "Mitigation"],
    ["Macro F1 ≥ 0.75 may not be reachable on 145 fine-grained classes without domain-tuned modeling",
     "Multi-pronged Week 5–6 plan: segment-only inputs + hierarchy-aware multi-task transformer + long-tail loss + retrieval augmentation. Each layer adds incrementally."],
    ["GPU compute limits on Colab free tier may bottleneck DeBERTa fine-tuning",
     "Outputs (weights, embeddings) downloaded back to local CPU for inference; demo remains fully offline."],
    ["Conglomerate label noise is fundamentally unresolvable",
     "If true, we present the honest ceiling with full diagnostic evidence rather than chase an illusion. Macro F1 in the 70–75% range honestly evaluated is still a stronger submission than a leaked 88%."],
    ["Team coordination across five members",
     "Per-person executable scripts (Week 4–5 task sheets) with no cross-dependencies; lead consolidates results."],
])
hrule()

# ─────────────────────────────────────────────────────────────────────────────
# 10. TIMELINE
# ─────────────────────────────────────────────────────────────────────────────
heading("10. Timeline and Next Steps", 1)
styled_table([
    ["Week", "Focus", "Owner"],
    ["Week 5 (May 10–16)",
     "Segment-only hypothesis test + hierarchical DeBERTa + Task 2 scaffolding",
     "Whole team, lead drives strategic experiments"],
    ["Week 6 (May 17–23)",
     "Long-tail loss (DB / DCAL), retrieval-augmented classifier, Task 2 baseline",
     "Lead + 2"],
    ["Week 7 (May 24–30)",
     "Ensemble best Task 1 model with Task 2 cross-constraint; demo cleanup with calibrated confidence",
     "Lead + 2"],
    ["Week 8 (May 31 – Jun 6)",
     "Error analysis, final write-up, reproducibility verification, presentation rehearsal",
     "Whole team"],
])
hrule()

# ─────────────────────────────────────────────────────────────────────────────
# 11. DELIVERABLES
# ─────────────────────────────────────────────────────────────────────────────
heading("11. What We Will Deliver", 1)
bullet("Working Task 1 classifier with Macro F1 ≥ 0.75 honestly evaluated, or a documented best-effort below 0.75 with full diagnostic evidence of the ceiling.")
bullet("Working Task 2 classifier that uses the Task 1 → Task 2 deterministic mapping as a hard constraint.")
bullet("Interactive demo running locally on port 5003, with calibrated probabilities and an honest top-3 panel.")
bullet("Audit document (CASCADE_AUDIT.md) detailing the leakage we caught, the splits, the iterations, and the methodology.")
bullet("Project journey document (PROJECT_JOURNEY.md) telling the full story.")
bullet("Reproducible training scripts for every numbered version (V1 through V20+), with per-run training_summary.json artifacts.")
bullet("Per-week team-task sheets showing how the work was distributed across the team.")
bullet("Final presentation suitable for Morningstar RED team review.")
hrule()

# ─────────────────────────────────────────────────────────────────────────────
# 12. CLOSING
# ─────────────────────────────────────────────────────────────────────────────
heading("12. Closing Statement", 1)
para("Five weeks in, we have learned that the difference between a credible 88.90% and an honest 69.09% on "
     "this dataset is methodology, not modeling. Catching that distinction ourselves — rather than presenting "
     "an inflated number to Morningstar and being asked the obvious follow-up question — is the most important "
     "thing we have done so far. The remaining weeks are about closing the gap honestly, exploiting structural "
     "signal (the GECS hierarchy, the official taxonomy text, the Task 1 → Task 2 constraint) that other teams "
     "will not have spent the time to extract.")
para("The number we deliver will be real. The work to get there is documented end-to-end.", bold=True)

doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot.add_run("Prepared by Group 4 · DePaul University · MGT 599 · Q2 2026 · May 10, 2026")
r.font.size = Pt(9)
r.italic = True

doc.save(OUT)
print(f"Wrote {OUT}")
